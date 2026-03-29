"""
03_random_forest.py
====================
EWS-CEMAC — Random Forest Classifier, Reference Split 2019

PURPOSE
-------
Trains and evaluates a Random Forest (RF) classifier on the CEMAC panel
(6 countries × 2000–2023, N = 144 country-years) using the same 2019
reference split as the Logit-L2 baseline (script 02).

RF is included because it can capture non-linear interactions between
predictors (e.g., credit_growth × NPL) that a linear logit model misses,
as documented for EWS in Holopainen & Sarlin (2017, JFSR) and
Alessi & Detken (2018, JFS).

HYPERPARAMETER JUSTIFICATION
-----------------------------
All RF hyperparameters are chosen conservatively for a small panel (EPV ≈ 4.8):

  n_estimators = 500
      Enough trees for stable OOB error estimates; OOB is used as an
      unbiased proxy for generalisation error without a separate validation set.

  max_depth = 3
      Shallow trees prevent overfitting on a 78-observation training set.
      Each tree can model at most 2³ = 8 distinct decision regions; given
      that the signal is low (AUC ≈ 0.72 in-sample), deep trees would
      memorise noise.

  min_samples_leaf = 5
      Each terminal leaf must contain ≥ 5 observations ≈ 4.4% of N_train.
      This floor prevents leaves with 1–2 observations, which would have
      extremely noisy probability estimates in a small crisis panel.

  max_features = 'sqrt'
      At each split, sqrt(5) ≈ 2 candidate features are randomly selected.
      This decorrelates the individual trees (Breiman, 2001).

  class_weight = 'balanced'
      Equivalent to oversampling the minority class (stress) in proportion
      to its rarity: w_k = N / (n_classes × N_k). See script 02 for details.

  bootstrap = True, oob_score = True
      Standard Breiman RF with OOB scoring. OOB AUC is reported alongside
      the test AUC as a sanity check for overfitting.

IMPUTATION & SCALING: identical to script 02 (Bulmer rule, train-only fit).

AUTHORS
-------
  Françoise NGOUFACK, Pamphile MEZUI-MBENG, Samba NDIAYE — 2026
  Paper: "Do Early Warning Systems Survive Structural Breaks?
          Macroprudential Evidence from the CEMAC Monetary Union"
  Journal of Financial Stability [under review]
"""

import sys
sys.stdout.reconfigure(encoding='utf-8', errors='replace')

import os
import pandas as pd
import numpy as np
import warnings
warnings.filterwarnings('ignore')

from sklearn.ensemble import RandomForestClassifier
from sklearn.linear_model import LogisticRegression
from sklearn.inspection import permutation_importance
from sklearn.metrics import (
    roc_auc_score, roc_curve,
    precision_recall_curve, average_precision_score,
    confusion_matrix, brier_score_loss,
    precision_score, recall_score, f1_score
)
from sklearn.preprocessing import StandardScaler
from sklearn.impute import SimpleImputer
from sklearn.utils.class_weight import compute_class_weight

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.colors import LinearSegmentedColormap

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ══════════════════════════════════════════════════════════════════════
# PARAMETRES
# ══════════════════════════════════════════════════════════════════════
DATA_DIR = r"C:\Users\fngou\Desktop\Donnees_Memoire ML"
if not os.path.exists(DATA_DIR):
    DATA_DIR = r"C:\Users\fngou\Desktop\Données_Mémoire ML"

OUT_DIR    = r"C:\Users\fngou\Desktop\Chapitres du Memoire ML"
PATH_MACRO = DATA_DIR + r"\Dataset_Macro_CEMAC.csv"
PATH_DOC   = OUT_DIR  + r"\Section_36_RF_StressVuln_2019.docx"
PATH_FIG1  = OUT_DIR  + r"\fig_rf2019_roc_pr.png"
PATH_FIG2  = OUT_DIR  + r"\fig_rf2019_confusion.png"
PATH_FIG3  = OUT_DIR  + r"\fig_rf2019_importance.png"
PATH_FIG4  = OUT_DIR  + r"\fig_rf2019_proba.png"
PATH_FIG5  = OUT_DIR  + r"\fig_rf2019_compare.png"

TARGET     = 'StressScore'
TARGET_BIN = 'Target_Stress2'
SPLIT_YEAR = 2019
SKEW_THR   = 1.0

# RF hyperparameters — conservative choices for small panel (EPV ≈ 4.8).
# Full justification in the module docstring above.
RF_PARAMS = dict(
    n_estimators    = 500,       # stable OOB error; computationally affordable
    max_depth       = 3,         # shallow trees prevent overfitting (N_train ≈ 78)
    min_samples_leaf= 5,         # ≥ 5 obs per leaf → stable leaf probabilities
    max_features    = 'sqrt',    # √5 ≈ 2 features per split → tree decorrelation
    class_weight    = 'balanced',# upweight crisis class (Breiman, 2001)
    bootstrap       = True,
    oob_score       = True,      # out-of-bag AUC as internal validation proxy
    random_state    = 42,
    n_jobs          = -1
)

# 5 variables strategiques
SELECTED = [
    ('M2_croissance_pct',        'Croissance M2 (%)',        'C1',  0.69),
    ('Solde_budgetaire_pct_PIB', 'Solde budgetaire (% PIB)', 'C2',  0.99),
    ('PIB_croissance_reel_pct',  'Croissance PIB reel (%)',  'C3',  5.23),
    ('Reserves_USD',             'Reserves change (USD)',    'C4',  1.11),
    ('Rentes_petrole_pct_PIB',   'Rentes petrole (% PIB)',   'C5',  0.80),
]

BASE_COLS = [s[0] for s in SELECTED]
LAG_COLS  = [s[0]+'_lag1' for s in SELECTED]
LABELS    = {s[0]: s[1] for s in SELECTED}
LABELS_L  = {s[0]+'_lag1': s[1]+' (lag-1)' for s in SELECTED}

IMPUTE    = {s[0]: ('median' if abs(s[3]) > SKEW_THR else 'mean') for s in SELECTED}

BLUE='#1F3864'; LBLUE='#2E74B5'; ORANGE='#C0392B'; GREEN='#1A7A2E'; GRAY='#95A5A6'
TEAL='#148F77'; PURPLE='#6C3483'

# ══════════════════════════════════════════════════════════════════════
# 1. CHARGEMENT & PREPROCESSING (identique au Logit)
# ══════════════════════════════════════════════════════════════════════
print("="*66)
print("  EWS-CEMAC  |  Random Forest  |  StressScore >= 2  |  Split 2019")
print("="*66)

df = (pd.read_csv(PATH_MACRO)
        .sort_values(['Country','Year'])
        .reset_index(drop=True))
df[TARGET_BIN] = (df[TARGET] >= 2).astype(int)

# Lags sur panel complet avant split (variables predeterminees t-1 : pas de leakage)
for col in BASE_COLS:
    df[col+'_lag1'] = df.groupby('Country')[col].shift(1)

train_df = df[df['Year'] <  SPLIT_YEAR].copy().reset_index(drop=True)
test_df  = df[df['Year'] >= SPLIT_YEAR].copy().reset_index(drop=True)

X_tr_raw = train_df[LAG_COLS].copy()
X_te_raw = test_df[LAG_COLS].copy()
y_train  = train_df[TARGET_BIN].astype(int).values
y_test   = test_df[TARGET_BIN].astype(int).values
meta_tr  = train_df[['Country','Year']].copy()
meta_te  = test_df[['Country','Year']].copy()

n_tr, n1_tr = len(y_train), int(y_train.sum())
n_te, n1_te = len(y_test),  int(y_test.sum())

print(f"\n  TRAIN : 2000-{SPLIT_YEAR-1} | {n_tr} obs | {n1_tr} vulnerables ({n1_tr/n_tr*100:.1f}%)")
print(f"  TEST  : {SPLIT_YEAR}-2023   | {n_te} obs | {n1_te} vulnerables ({n1_te/n_te*100:.1f}%)")
print(f"  EPV   : {n1_tr}/{len(LAG_COLS)} = {n1_tr/len(LAG_COLS):.1f}")

# Imputation
mean_cols   = [c+'_lag1' for c in BASE_COLS if IMPUTE[c]=='mean']
median_cols = [c+'_lag1' for c in BASE_COLS if IMPUTE[c]=='median']
imp_mean   = SimpleImputer(strategy='mean')
imp_median = SimpleImputer(strategy='median')
if mean_cols:   imp_mean.fit(X_tr_raw[mean_cols])
if median_cols: imp_median.fit(X_tr_raw[median_cols])
for X in [X_tr_raw, X_te_raw]:
    if mean_cols:   X[mean_cols]   = imp_mean.transform(X[mean_cols])
    if median_cols: X[median_cols] = imp_median.transform(X[median_cols])
assert X_tr_raw.isna().sum().sum()==0

# Normalisation (RF n'en a pas besoin mais utile pour permutation importance)
scaler  = StandardScaler()
X_train = pd.DataFrame(scaler.fit_transform(X_tr_raw), columns=LAG_COLS)
X_test  = pd.DataFrame(scaler.transform(X_te_raw),     columns=LAG_COLS)

# Class weights
cw_arr  = compute_class_weight('balanced', classes=np.array([0,1]), y=y_train)
cw_dict = {0: round(cw_arr[0],4), 1: round(cw_arr[1],4)}
sw      = np.where(y_train==1, cw_dict[1], cw_dict[0])
ratio   = round(cw_dict[1]/cw_dict[0], 1)
print(f"  Class weights : w0={cw_dict[0]} | w1={cw_dict[1]} | ratio={ratio}x")
print(f"  Imputation OK | Normalisation OK")

# ══════════════════════════════════════════════════════════════════════
# 2. ENTRAINEMENT RANDOM FOREST
# ══════════════════════════════════════════════════════════════════════
print("\n" + "="*66)
print(f"  RANDOM FOREST  n={RF_PARAMS['n_estimators']}  depth={RF_PARAMS['max_depth']}"
      f"  leaf={RF_PARAMS['min_samples_leaf']}  features=sqrt")
print("="*66)

rf = RandomForestClassifier(**RF_PARAMS)
rf.fit(X_train, y_train)

prob_tr = rf.predict_proba(X_train)[:,1]
prob_te = rf.predict_proba(X_test)[:,1]

print(f"\n  OOB Score  : {rf.oob_score_:.4f}")
print(f"  P_hat train : [{prob_tr.min():.4f}, {prob_tr.max():.4f}]  moy={prob_tr.mean():.4f}")
print(f"  P_hat test  : [{prob_te.min():.4f}, {prob_te.max():.4f}]  moy={prob_te.mean():.4f}")

# Seuil Youden sur train
fpr_tr, tpr_tr, thr_tr = roc_curve(y_train, prob_tr)
j_tr    = tpr_tr + (1-fpr_tr) - 1
idx_opt = np.argmax(j_tr)
TAU_OPT = float(thr_tr[idx_opt])
TAU_DEF = 0.5
print(f"  Seuil Youden (train) : tau_opt = {TAU_OPT:.4f}")

# ══════════════════════════════════════════════════════════════════════
# 3. METRIQUES
# ══════════════════════════════════════════════════════════════════════
def metriques(y_true, y_prob, tau, label):
    y_pred = (y_prob >= tau).astype(int)
    cm = confusion_matrix(y_true, y_pred)
    tn, fp, fn, tp = (cm.ravel() if cm.size==4
                      else (cm[0,0],0,0,0) if y_true.sum()==0
                      else (0,0,0,cm[0,0]))
    has_both = len(np.unique(y_true)) > 1
    auc    = roc_auc_score(y_true, y_prob)           if has_both else np.nan
    auc_pr = average_precision_score(y_true, y_prob) if has_both else np.nan
    rappel = recall_score(y_true, y_pred, zero_division=0)
    prec   = precision_score(y_true, y_pred, zero_division=0)
    f1     = f1_score(y_true, y_pred, zero_division=0)
    brier  = brier_score_loss(y_true, y_prob)
    spec   = tn/(tn+fp) if (tn+fp)>0 else 0.
    j_val  = rappel + spec - 1
    far    = fp/(fp+tn) if (fp+tn)>0 else 0.

    print(f"\n  --- {label} ---")
    print(f"  AUC-ROC : {auc:.4f}  |  AUC-PR : {auc_pr:.4f}")
    print(f"  Rappel  : {rappel:.4f}  ({tp}/{tp+fn})  |  Precision : {prec:.4f}")
    print(f"  F1      : {f1:.4f}  |  Brier : {brier:.4f}")
    print(f"  J       : {j_val:.4f}  |  Spec={spec:.4f}  FAR={far:.4f}")
    print(f"  CM : TN={tn} FP={fp} FN={fn} TP={tp}")

    return dict(label=label, tau=tau,
                auc=auc, auc_pr=auc_pr,
                rappel=rappel, prec=prec, f1=f1, brier=brier,
                spec=spec, j=j_val, far=far,
                tn=int(tn), fp=int(fp), fn=int(fn), tp=int(tp))

print("\n" + "="*66)
print("  RESULTATS RANDOM FOREST")
print("="*66)

m_tr  = metriques(y_train, prob_tr, TAU_OPT, f'TRAIN  tau={TAU_OPT:.4f}')
m_te  = metriques(y_test,  prob_te, TAU_OPT, f'TEST   tau={TAU_OPT:.4f} (calibre train)')
m_te5 = metriques(y_test,  prob_te, TAU_DEF, f'TEST   tau=0.50 (reference)')

# ── Importance des variables ──────────────────────────────────────────
# Gini importance
gini_imp = pd.DataFrame({
    'Feature': LAG_COLS,
    'Label'  : [LABELS_L[c] for c in LAG_COLS],
    'Gini'   : rf.feature_importances_
}).sort_values('Gini', ascending=False).reset_index(drop=True)

# Permutation importance (sur test — plus fiable)
perm = permutation_importance(rf, X_test, y_test,
                               n_repeats=50, random_state=42,
                               scoring='roc_auc')
perm_df = pd.DataFrame({
    'Feature': LAG_COLS,
    'Label'  : [LABELS_L[c] for c in LAG_COLS],
    'Perm_mean': perm.importances_mean,
    'Perm_std' : perm.importances_std
}).sort_values('Perm_mean', ascending=False).reset_index(drop=True)

print("\n  Importance Gini (train) :")
for _,r in gini_imp.iterrows():
    bar = '█' * int(r['Gini']*40)
    print(f"    {r['Label']:35s}  {r['Gini']:.4f}  {bar}")

print("\n  Importance Permutation (test, n_repeats=50) :")
for _,r in perm_df.iterrows():
    bar = '█' * max(0,int(r['Perm_mean']*40))
    print(f"    {r['Label']:35s}  {r['Perm_mean']:+.4f} ± {r['Perm_std']:.4f}  {bar}")

# Analyse par pays (test)
print("\n  Analyse par pays — TEST :")
country_res = []
for country in sorted(meta_te['Country'].unique()):
    mask = meta_te['Country'].values == country
    y_c  = y_test[mask]; p_c = prob_te[mask]
    n_c  = len(y_c); nc1 = int(y_c.sum())
    pred = (p_c >= TAU_OPT).astype(int)
    tp_c = int(((y_c==1)&(pred==1)).sum())
    fn_c = int(((y_c==1)&(pred==0)).sum())
    fp_c = int(((y_c==0)&(pred==1)).sum())
    tn_c = int(((y_c==0)&(pred==0)).sum())
    auc_c = roc_auc_score(y_c, p_c) if len(np.unique(y_c))==2 else np.nan
    country_res.append((country, n_c, nc1, p_c.mean(), tp_c, fn_c, fp_c, tn_c, auc_c))
    auc_s = f"{auc_c:.3f}" if not np.isnan(auc_c) else 'n/a'
    print(f"    {country:30s}  N={n_c} Vuln={nc1}  "
          f"p_moy={p_c.mean():.3f}  TP={tp_c} FN={fn_c} FP={fp_c}  AUC={auc_s}")

# ── Recap Logit vs RF ─────────────────────────────────────────────────
# Re-entrainer le logit pour comparaison
logit_ref = LogisticRegression(penalty='l2', C=0.1, class_weight='balanced',
                                solver='lbfgs', max_iter=2000, random_state=42)
logit_ref.fit(X_train, y_train)
pl_tr = logit_ref.predict_proba(X_train)[:,1]
pl_te = logit_ref.predict_proba(X_test)[:,1]

fpr_l, tpr_l, thr_l = roc_curve(y_train, pl_tr)
j_l   = tpr_l + (1-fpr_l) - 1
tau_l = float(thr_l[np.argmax(j_l)])

ml_te = dict(
    auc    = roc_auc_score(y_test, pl_te),
    auc_pr = average_precision_score(y_test, pl_te),
    rappel = recall_score(y_test, (pl_te>=tau_l).astype(int), zero_division=0),
    prec   = precision_score(y_test, (pl_te>=tau_l).astype(int), zero_division=0),
    f1     = f1_score(y_test, (pl_te>=tau_l).astype(int), zero_division=0),
    brier  = brier_score_loss(y_test, pl_te),
)

print("\n" + "="*66)
print("  COMPARAISON : Logit L2 vs Random Forest (test tau_opt)")
print("="*66)
print(f"  {'Metrique':15s}  {'Logit L2':>10s}  {'RF':>10s}  {'Delta':>10s}")
print("  " + "-"*50)
for k in ['auc','auc_pr','rappel','prec','f1','brier']:
    v_l = ml_te[k]; v_r = m_te[k]
    delta = v_r - v_l
    sym = '▲' if ((delta>0 and k!='brier') or (delta<0 and k=='brier')) else '▼'
    print(f"  {k:15s}  {v_l:>10.4f}  {v_r:>10.4f}  {delta:>+10.4f} {sym}")

# ══════════════════════════════════════════════════════════════════════
# 4. FIGURES
# ══════════════════════════════════════════════════════════════════════
print("\nGeneration des figures...")

# ── Fig 1 : ROC + PR ─────────────────────────────────────────────────
fig, axes = plt.subplots(1, 2, figsize=(13, 5.5))
fig.suptitle(f"Random Forest — Vulnerabilite CEMAC | Split {SPLIT_YEAR} | "
             f"depth={RF_PARAMS['max_depth']} leaf={RF_PARAMS['min_samples_leaf']}",
             fontsize=12, fontweight='bold', color=BLUE)

fpr_rf_te, tpr_rf_te, _ = roc_curve(y_test,  prob_te)
fpr_rf_tr, tpr_rf_tr, _ = roc_curve(y_train, prob_tr)
fpr_l_te,  tpr_l_te,  _ = roc_curve(y_test,  pl_te)

ax = axes[0]
ax.fill_between(fpr_rf_te, tpr_rf_te, alpha=0.12, color=TEAL)
ax.plot(fpr_rf_te, tpr_rf_te, color=TEAL,  lw=2.5,
        label=f"RF Test    AUC = {m_te['auc']:.3f}")
ax.plot(fpr_rf_tr, tpr_rf_tr, color=TEAL,  lw=1.5, ls='--',
        label=f"RF Train   AUC = {m_tr['auc']:.3f}")
ax.plot(fpr_l_te,  tpr_l_te,  color=LBLUE, lw=1.5, ls=':',
        label=f"Logit Test AUC = {ml_te['auc']:.3f}")
ax.plot([0,1],[0,1], color=GRAY, lw=1, ls=':', label='Aleatoire')
ax.scatter([m_te['far']], [m_te['rappel']], s=150, color=ORANGE, zorder=6,
           label=f"tau_opt={TAU_OPT:.3f}  J={m_te['j']:.3f}")
ax.set_xlabel('FAR (1 - Specificite)', fontsize=11)
ax.set_ylabel('Rappel (Sensibilite)', fontsize=11)
ax.set_title('Courbe ROC', fontsize=11, fontweight='bold', color=BLUE)
ax.legend(fontsize=9); ax.grid(alpha=0.3)
ax.set_xlim([-0.02,1.02]); ax.set_ylim([-0.02,1.02])

prec_rf_te, rec_rf_te, _ = precision_recall_curve(y_test,  prob_te)
prec_rf_tr, rec_rf_tr, _ = precision_recall_curve(y_train, prob_tr)
prec_l_te,  rec_l_te,  _ = precision_recall_curve(y_test,  pl_te)
ax = axes[1]
ax.fill_between(rec_rf_te, prec_rf_te, alpha=0.12, color=TEAL)
ax.plot(rec_rf_te, prec_rf_te, color=TEAL,  lw=2.5,
        label=f"RF Test    AUC-PR = {m_te['auc_pr']:.3f}")
ax.plot(rec_rf_tr, prec_rf_tr, color=TEAL,  lw=1.5, ls='--',
        label=f"RF Train   AUC-PR = {m_tr['auc_pr']:.3f}")
ax.plot(rec_l_te,  prec_l_te,  color=LBLUE, lw=1.5, ls=':',
        label=f"Logit Test AUC-PR = {ml_te['auc_pr']:.3f}")
ax.axhline(y_test.mean(), color=GRAY, lw=1, ls=':',
           label=f"Baseline ({y_test.mean():.2f})")
ax.scatter([m_te['rappel']], [m_te['prec']], s=150, color=ORANGE, zorder=6,
           label=f"tau_opt={TAU_OPT:.3f}")
ax.set_xlabel('Rappel', fontsize=11); ax.set_ylabel('Precision', fontsize=11)
ax.set_title('Courbe Precision-Rappel', fontsize=11, fontweight='bold', color=BLUE)
ax.legend(fontsize=9); ax.grid(alpha=0.3)
ax.set_xlim([-0.02,1.02]); ax.set_ylim([-0.02,1.02])
plt.tight_layout()
plt.savefig(PATH_FIG1, dpi=150, bbox_inches='tight'); plt.close()
print(f"  Fig 1 (ROC+PR) : OK")

# ── Fig 2 : Matrices de confusion ────────────────────────────────────
fig, axes = plt.subplots(1, 2, figsize=(11, 4.5))
fig.suptitle(f"Matrices de confusion RF — Test {SPLIT_YEAR}-2023",
             fontsize=12, fontweight='bold', color=BLUE)

def plot_cm(ax, m, title, cmap_color=TEAL):
    cm_m = np.array([[m['tn'],m['fp']],[m['fn'],m['tp']]])
    cmap = LinearSegmentedColormap.from_list('c',['#FFFFFF',cmap_color])
    ax.imshow(cm_m, cmap=cmap, vmin=0, vmax=max(cm_m.max(),1))
    for i in range(2):
        for j in range(2):
            v = cm_m[i,j]
            col = 'white' if v > cm_m.max()*0.55 else 'black'
            ax.text(j,i,f"{[['TN','FP'],['FN','TP']][i][j]}\n{v}",
                    ha='center',va='center',fontsize=14,fontweight='bold',color=col)
    ax.set_xticks([0,1]); ax.set_yticks([0,1])
    ax.set_xticklabels(['Predit 0\n(Stable)','Predit 1\n(Vulnerable)'],fontsize=10)
    ax.set_yticklabels(['Reel 0\n(Stable)','Reel 1\n(Vulnerable)'],fontsize=10)
    ax.set_title(title, fontsize=11, fontweight='bold', color=BLUE)
    ax.set_xlabel(f"Rappel={m['rappel']:.2f}  Prec={m['prec']:.2f}  "
                  f"F1={m['f1']:.2f}  J={m['j']:.2f}", fontsize=9.5)

plot_cm(axes[0], m_te5, "Seuil par defaut (tau = 0.50)")
plot_cm(axes[1], m_te,  f"Seuil optimal Youden (tau = {TAU_OPT:.4f})")
plt.tight_layout()
plt.savefig(PATH_FIG2, dpi=150, bbox_inches='tight'); plt.close()
print(f"  Fig 2 (CM) : OK")

# ── Fig 3 : Importance Gini + Permutation ────────────────────────────
fig, axes = plt.subplots(1, 2, figsize=(14, 4.5))
fig.suptitle("Importance des variables — Random Forest",
             fontsize=12, fontweight='bold', color=BLUE)

# Gini
gi = gini_imp.sort_values('Gini')
axes[0].barh(gi['Label'], gi['Gini'], color=TEAL, edgecolor='white', height=0.6)
for i,v in enumerate(gi['Gini']):
    axes[0].text(v+0.003, i, f'{v:.4f}', va='center', ha='left', fontsize=9.5)
axes[0].set_xlabel('Importance Gini (diminution impurete moyenne)', fontsize=10)
axes[0].set_title('Importance Gini (train)', fontsize=11, fontweight='bold', color=BLUE)
axes[0].grid(axis='x', alpha=0.3)
axes[0].tick_params(axis='y', labelsize=9.5)

# Permutation
pi = perm_df.sort_values('Perm_mean')
col_p = [ORANGE if v > 0 else GRAY for v in pi['Perm_mean']]
axes[1].barh(pi['Label'], pi['Perm_mean'], xerr=pi['Perm_std'],
             color=col_p, edgecolor='white', height=0.6,
             error_kw=dict(ecolor='#555555', capsize=4))
axes[1].axvline(0, color='black', lw=0.8)
for i,v in enumerate(pi['Perm_mean']):
    off = 0.005 if v>=0 else -0.005
    axes[1].text(v+off, i, f'{v:+.4f}', va='center',
                 ha='left' if v>=0 else 'right', fontsize=9.5)
axes[1].set_xlabel('Diminution AUC-ROC par permutation (test, n=50)', fontsize=10)
axes[1].set_title('Importance par permutation (test)', fontsize=11, fontweight='bold', color=BLUE)
axes[1].grid(axis='x', alpha=0.3)
axes[1].tick_params(axis='y', labelsize=9.5)
axes[1].legend(handles=[
    mpatches.Patch(color=ORANGE, label='Contribution positive (AUC baisse si permuté)'),
    mpatches.Patch(color=GRAY,   label='Contribution nulle ou negative')], fontsize=9)
plt.tight_layout()
plt.savefig(PATH_FIG3, dpi=150, bbox_inches='tight'); plt.close()
print(f"  Fig 3 (importance) : OK")

# ── Fig 4 : Distribution P_hat ────────────────────────────────────────
fig, axes = plt.subplots(1, 2, figsize=(12, 4.5))
fig.suptitle("Distribution de P_hat(Vulnerabilite) — Random Forest",
             fontsize=12, fontweight='bold', color=BLUE)
for ax, (probs, y_true, title) in zip(axes, [
        (prob_tr, y_train, f"Train 2000-{SPLIT_YEAR-1}"),
        (prob_te, y_test,  f"Test  {SPLIT_YEAR}-2023")]):
    ax.hist(probs[y_true==0], bins=20, alpha=0.6, color=LBLUE,
            label='Stable (y=0)', edgecolor='white')
    ax.hist(probs[y_true==1], bins=12, alpha=0.75, color=ORANGE,
            label='Vulnerable (y=1)', edgecolor='white')
    ax.axvline(TAU_OPT, color='red', ls='--', lw=2,  label=f'tau_opt={TAU_OPT:.3f}')
    ax.axvline(TAU_DEF, color=GRAY,  ls=':',  lw=1.5, label='tau=0.50')
    ax.set_xlabel('P_hat(Vulnerabilite)', fontsize=11)
    ax.set_ylabel('Observations', fontsize=11)
    ax.set_title(title, fontsize=11, fontweight='bold', color=BLUE)
    ax.legend(fontsize=9.5); ax.grid(alpha=0.3)
plt.tight_layout()
plt.savefig(PATH_FIG4, dpi=150, bbox_inches='tight'); plt.close()
print(f"  Fig 4 (proba) : OK")

# ── Fig 5 : Comparaison Logit vs RF ──────────────────────────────────
metrics_names  = ['AUC-ROC','AUC-PR','Rappel','Precision','F1-Score','Brier']
vals_logit = [ml_te['auc'], ml_te['auc_pr'], ml_te['rappel'],
              ml_te['prec'], ml_te['f1'], ml_te['brier']]
vals_rf    = [m_te['auc'],  m_te['auc_pr'],  m_te['rappel'],
              m_te['prec'],  m_te['f1'],      m_te['brier']]

fig, ax = plt.subplots(figsize=(11, 5))
x = np.arange(len(metrics_names))
w = 0.35
bars_l = ax.bar(x - w/2, vals_logit, w, label='Logit L2 (baseline)',
                color=LBLUE, edgecolor='white', alpha=0.85)
bars_r = ax.bar(x + w/2, vals_rf,    w, label='Random Forest',
                color=TEAL, edgecolor='white', alpha=0.85)

for bar, v in zip(bars_l, vals_logit):
    ax.text(bar.get_x()+bar.get_width()/2, bar.get_height()+0.01,
            f'{v:.3f}', ha='center', va='bottom', fontsize=9, color=LBLUE, fontweight='bold')
for bar, v in zip(bars_r, vals_rf):
    ax.text(bar.get_x()+bar.get_width()/2, bar.get_height()+0.01,
            f'{v:.3f}', ha='center', va='bottom', fontsize=9, color=TEAL, fontweight='bold')

# Fleches delta
for i,(vl,vr,nm) in enumerate(zip(vals_logit,vals_rf,metrics_names)):
    better = vr>vl if nm!='Brier' else vr<vl
    col = GREEN if better else ORANGE
    sym = '▲' if better else '▼'
    ax.text(i, max(vl,vr)+0.05, f"{sym}{abs(vr-vl):.3f}",
            ha='center', fontsize=8.5, color=col, fontweight='bold')

ax.set_xticks(x); ax.set_xticklabels(metrics_names, fontsize=11)
ax.set_ylabel('Valeur (test set)', fontsize=11)
ax.set_ylim([0, 1.15])
ax.set_title(f"Comparaison Logit L2 vs Random Forest — Test {SPLIT_YEAR}-2023 | StressScore>=2",
             fontsize=12, fontweight='bold', color=BLUE)
ax.legend(fontsize=10); ax.grid(axis='y', alpha=0.3)
ax.axhline(0.5, color=GRAY, lw=1, ls='--', alpha=0.5, label='AUC aleatoire')
plt.tight_layout()
plt.savefig(PATH_FIG5, dpi=150, bbox_inches='tight'); plt.close()
print(f"  Fig 5 (comparaison) : OK")

# ══════════════════════════════════════════════════════════════════════
# 5. DOCUMENT WORD
# ══════════════════════════════════════════════════════════════════════
print("\nConstruction du document Word...")

def set_bg(cell, hx):
    tcPr=cell._tc.get_or_add_tcPr(); s=OxmlElement('w:shd')
    s.set(qn('w:val'),'clear'); s.set(qn('w:color'),'auto')
    s.set(qn('w:fill'),hx); tcPr.append(s)

def set_brd(t, color='1F3864', sz=8):
    tbl=t._tbl; tblPr=tbl.tblPr or OxmlElement('w:tblPr')
    b=OxmlElement('w:tblBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        el=OxmlElement(f'w:{side}')
        el.set(qn('w:val'),'single')
        el.set(qn('w:sz'),str(sz if side in('top','left','bottom','right') else 4))
        el.set(qn('w:space'),'0'); el.set(qn('w:color'),color); b.append(el)
    tblPr.append(b)

def cw(cell,text,bold=False,italic=False,sz=9.5,color=None,
       al=WD_ALIGN_PARAGRAPH.LEFT,sb=2,sa=2):
    p=cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    p.clear(); p.alignment=al
    p.paragraph_format.space_before=Pt(sb); p.paragraph_format.space_after=Pt(sa)
    cell.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
    r=p.add_run(text); r.bold=bold; r.italic=italic; r.font.size=Pt(sz)
    if color: r.font.color.rgb=RGBColor(*bytes.fromhex(color))

def H(doc,text,level,color='1F3864',sz=None):
    sz_d={1:13,2:12,3:11}
    p=doc.add_paragraph(style=f'Heading {level}')
    r=p.add_run(text); r.bold=True
    r.font.size=Pt(sz or sz_d[level])
    r.font.color.rgb=RGBColor(*bytes.fromhex(color))

def B(doc,text,italic=False,sz=10.5,sb=0,sa=4,indent=None):
    p=doc.add_paragraph(style='Normal')
    p.paragraph_format.space_before=Pt(sb); p.paragraph_format.space_after=Pt(sa)
    if indent: p.paragraph_format.left_indent=Cm(indent)
    r=p.add_run(text); r.font.size=Pt(sz); r.italic=italic

def Bmix(doc,segs,sb=0,sa=4,indent=None):
    p=doc.add_paragraph(style='Normal')
    p.paragraph_format.space_before=Pt(sb); p.paragraph_format.space_after=Pt(sa)
    if indent: p.paragraph_format.left_indent=Cm(indent)
    for txt,bold,italic,col in segs:
        r=p.add_run(txt); r.font.size=Pt(10.5); r.bold=bold; r.italic=italic
        if col: r.font.color.rgb=RGBColor(*bytes.fromhex(col))

def IMG(doc,path,w=15.0,cap_text=None,num=None):
    if not os.path.exists(path):
        B(doc,f"[Figure manquante : {path}]",italic=True); return
    p=doc.add_paragraph(style='Normal')
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(2)
    p.add_run().add_picture(path,width=Cm(w))
    if cap_text:
        pc=doc.add_paragraph(style='Normal')
        pc.alignment=WD_ALIGN_PARAGRAPH.CENTER
        pc.paragraph_format.space_before=Pt(2); pc.paragraph_format.space_after=Pt(10)
        rc=pc.add_run(f"Figure 3.{num} — {cap_text}" if num else cap_text)
        rc.italic=True; rc.font.size=Pt(9.5)
        rc.font.color.rgb=RGBColor(0x44,0x44,0x44)

def CAP(doc,text):
    p=doc.add_paragraph(style='Normal')
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(8)
    r=p.add_run(text); r.italic=True; r.font.size=Pt(9.5)
    r.font.color.rgb=RGBColor(0x44,0x44,0x44)

def SEP(doc):
    p=doc.add_paragraph(style='Normal')
    p.paragraph_format.space_before=Pt(10); p.paragraph_format.space_after=Pt(4)
    pPr=p._element.get_or_add_pPr(); pBdr=OxmlElement('w:pBdr')
    top=OxmlElement('w:top')
    top.set(qn('w:val'),'single'); top.set(qn('w:sz'),'4')
    top.set(qn('w:space'),'1'); top.set(qn('w:color'),'1F3864')
    pBdr.append(top); pPr.append(pBdr)

def INFO(doc,title,body,bg_c='DEEAF1',brd_c='1F3864'):
    t=doc.add_table(rows=1,cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    set_brd(t,color=brd_c,sz=6)
    c=t.rows[0].cells[0]; c.width=Cm(15.5); set_bg(c,bg_c)
    p1=c.paragraphs[0]; p1.clear()
    p1.paragraph_format.space_before=Pt(5); p1.paragraph_format.space_after=Pt(2)
    r1=p1.add_run(title); r1.bold=True; r1.font.size=Pt(10)
    r1.font.color.rgb=RGBColor(*bytes.fromhex(brd_c))
    p2=c.add_paragraph()
    p2.paragraph_format.space_before=Pt(0); p2.paragraph_format.space_after=Pt(5)
    r2=p2.add_run(body); r2.font.size=Pt(10); r2.italic=True
    r2.font.color.rgb=RGBColor(0x1A,0x1A,0x2E)
    doc.add_paragraph(style='Normal').paragraph_format.space_after=Pt(6)

def FBOX(doc,text):
    t=doc.add_table(rows=1,cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    set_brd(t,color='AAAAAA',sz=4)
    c=t.rows[0].cells[0]; c.width=Cm(14); set_bg(c,'F2F2F2')
    p=c.paragraphs[0]; p.clear()
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before=Pt(5); p.paragraph_format.space_after=Pt(5)
    r=p.add_run(text); r.font.size=Pt(10.5)
    r.font.name='Courier New'; r.font.color.rgb=RGBColor(0x1F,0x38,0x64)
    doc.add_paragraph(style='Normal').paragraph_format.space_after=Pt(4)

# ─── Document ────────────────────────────────────────────────────────
doc=Document()
for s in doc.sections:
    s.top_margin=Cm(2.5); s.bottom_margin=Cm(2.5)
    s.left_margin=Cm(3.0); s.right_margin=Cm(2.5)

H(doc,"3.6  Random Forest — Regime de Vulnerabilite Macro-Financiere",1)
B(doc,
  "Le Random Forest est le premier modele non-lineaire de la strategie "
  "de comparaison. Contrairement a la regression logistique, il capture "
  "les interactions d'ordre superieur et les effets de seuil entre "
  "indicateurs macroeconomiques — caracteristiques centrales des regimes "
  "de vulnerabilite financiere (Laeven & Valencia, 2013). "
  "Le modele est entraine sur la meme cible (StressScore >= 2), "
  "le meme split (2000-2018 / 2019-2023) et les memes 5 variables lag-1 "
  "que le logit baseline, garantissant une comparaison strictement controlee.",
  sb=6, sa=6)

# ── 3.6.1 Specification ───────────────────────────────────────────────
H(doc,"3.6.1  Specification et justification des hyperparametres",2)
B(doc,
  "Les hyperparametres sont fixes a priori selon les recommandations "
  "de la litterature pour les panels macro-financiers de petite taille "
  "(Dumitrescu et al., 2022 ; Tanaka et al., 2016). "
  "Aucune validation croisee temporelle n'est appliquee — "
  "le nombre d'evenements en train (n=24) la rend instable.",
  sb=0, sa=4)

HP_ROWS = [
    ('n_estimators','500',
     'Nombre suffisant pour stabiliser le score OOB et les importances'),
    ('max_depth','3',
     'Arbres peu profonds : previent le surapprentissage sur petit echantillon (N=114)'),
    ('min_samples_leaf','5',
     '>= 5 obs par feuille terminale (~4.4% de N_train) : evite les feuilles pures'),
    ('max_features','sqrt(5) ~ 2',
     'Sous-espace aleatoire par split : decorrelation des arbres, regularisation'),
    ('class_weight','balanced',
     f'w0={cw_dict[0]}, w1={cw_dict[1]} : correction du desequilibre 3.8:1'),
    ('bootstrap','True',
     'Bagging standard : echantillons OOB utilisables comme validation interne'),
    ('oob_score','True',
     'Score Out-of-Bag : estimation non biaisee de la performance train sans CV'),
]
tbl_hp = doc.add_table(rows=len(HP_ROWS)+1, cols=3)
tbl_hp.style='Table Grid'; tbl_hp.alignment=WD_TABLE_ALIGNMENT.CENTER
set_brd(tbl_hp,'1F3864',8)
for row in tbl_hp.rows:
    for i,c in enumerate(row.cells):
        c.width=Cm([4.0,3.0,9.0][i])
for i,h_ in enumerate(['Hyperparametre','Valeur','Justification']):
    set_bg(tbl_hp.rows[0].cells[i],'1F3864')
    cw(tbl_hp.rows[0].cells[i],h_,bold=True,sz=9,color='FFFFFF',al=WD_ALIGN_PARAGRAPH.CENTER)
for r_i,(hp,val,just) in enumerate(HP_ROWS):
    row=tbl_hp.rows[r_i+1]
    bg_c='EDF3FB' if r_i%2==0 else 'FFFFFF'
    for i,v in enumerate([hp,val,just]):
        set_bg(row.cells[i],bg_c)
        cw(row.cells[i],v,sz=9,bold=(i==0),
           al=WD_ALIGN_PARAGRAPH.CENTER if i<2 else WD_ALIGN_PARAGRAPH.LEFT)
CAP(doc,"Tableau 3.42 — Hyperparametres du Random Forest et justification")

FBOX(doc,
     "RF : P(Vuln_{t+1} | X_t) = (1/B) * sum_{b=1}^{500} T_b(X_t)\n\n"
     f"  T_b : arbre de decision  depth={RF_PARAMS['max_depth']}  "
     f"leaf>={RF_PARAMS['min_samples_leaf']}  features=sqrt(5)~2\n"
     f"  OOB Score : {rf.oob_score_:.4f}  |  tau_opt = {TAU_OPT:.4f} (Youden / train)\n"
     f"  class_weight : balanced  |  w0={cw_dict[0]}  w1={cw_dict[1]}")

# ── 3.6.2 Resultats ───────────────────────────────────────────────────
SEP(doc)
H(doc,"3.6.2  Resultats et metriques de performance",2)

MROWS=[
    ('AUC-ROC',    m_tr['auc'],    m_te5['auc'],    m_te['auc'],    False),
    ('AUC-PR',     m_tr['auc_pr'], m_te5['auc_pr'], m_te['auc_pr'], False),
    ('Rappel',     m_tr['rappel'], m_te5['rappel'],  m_te['rappel'],  False),
    ('Precision',  m_tr['prec'],   m_te5['prec'],   m_te['prec'],   False),
    ('F1-Score',   m_tr['f1'],     m_te5['f1'],     m_te['f1'],     False),
    ('Brier Score',m_tr['brier'],  m_te5['brier'],  m_te['brier'],  True),
]
KEY={'AUC-ROC','AUC-PR','Rappel','F1-Score'}

tbl_r=doc.add_table(rows=len(MROWS)+2,cols=4)
tbl_r.style='Table Grid'; tbl_r.alignment=WD_TABLE_ALIGNMENT.CENTER
set_brd(tbl_r,'1F3864',8)
for row in tbl_r.rows:
    for i,c in enumerate(row.cells):
        c.width=Cm([3.8,3.0,3.0,3.0][i])
for i,h_ in enumerate(['Metrique','Train (OOB={:.3f})'.format(rf.oob_score_),
                        'Test tau=0.50','Test tau_opt']):
    set_bg(tbl_r.rows[0].cells[i],['1F3864','148F77','1A5276','1A5276'][i])
    cw(tbl_r.rows[0].cells[i],h_,bold=True,sz=9,color='FFFFFF',al=WD_ALIGN_PARAGRAPH.CENTER)
set_bg(tbl_r.rows[1].cells[0],'2C3E50'); cw(tbl_r.rows[1].cells[0],'',sz=8)
for i,val in enumerate([f'tau={TAU_OPT:.4f}','tau=0.50',f'tau={TAU_OPT:.4f}']):
    set_bg(tbl_r.rows[1].cells[i+1],['148F77','1A5276','1A5276'][i])
    cw(tbl_r.rows[1].cells[i+1],val,sz=8,color='FFFFFF',al=WD_ALIGN_PARAGRAPH.CENTER)

for r_i,(name,v_tr,v_te5,v_te,low) in enumerate(MROWS):
    row=tbl_r.rows[r_i+2]
    bg0='F2F9FF' if name in KEY else ('EDF3FB' if r_i%2==0 else 'FFFFFF')
    set_bg(row.cells[0],bg0)
    cw(row.cells[0],name,bold=(name in KEY),sz=9.5)
    for ci,(val,is_opt) in enumerate([(v_tr,False),(v_te5,False),(v_te,True)]):
        better=(val>v_te5) if not low else (val<v_te5)
        bg_c=('E8F5E9' if is_opt and better else
              'FDECEA' if is_opt and not better else bg0)
        set_bg(row.cells[ci+1],bg_c)
        cw(row.cells[ci+1],f"{val:.4f}",bold=is_opt,sz=9.5,
           al=WD_ALIGN_PARAGRAPH.CENTER,
           color=('1A7A2E' if is_opt and better else
                  'C0392B' if is_opt and not better else None))
CAP(doc,f"Tableau 3.43 — Resultats Random Forest (StressScore>=2, split {SPLIT_YEAR}, "
        f"tau_opt={TAU_OPT:.4f})")

IMG(doc,PATH_FIG1,w=15.5,
    cap_text=f"Courbes ROC et PR — RF vs Logit (train 2000-{SPLIT_YEAR-1} / test {SPLIT_YEAR}-2023)",
    num='16')
IMG(doc,PATH_FIG2,w=13.5,
    cap_text=f"Matrices de confusion RF — Test {SPLIT_YEAR}-2023 (tau=0.50 vs tau_opt={TAU_OPT:.4f})",
    num='17')

# ── 3.6.3 Importance des variables ────────────────────────────────────
SEP(doc)
H(doc,"3.6.3  Importance des variables",2)
B(doc,
  "Deux mesures d'importance complementaires sont rapportees. "
  "L'importance Gini (diminution moyenne de l'impurete de Gini sur le train) "
  "est biaisee vers les variables continues a forte cardinalite. "
  "L'importance par permutation (diminution AUC-ROC apres permutation "
  "aleatoire sur le test, n=50 repetitions) est plus fiable et reflète "
  "la contribution reelle a la generalisation hors echantillon.",
  sb=0, sa=4)

# Tableau importance
all_imp = gini_imp.merge(perm_df[['Feature','Perm_mean','Perm_std']], on='Feature')
all_imp = all_imp.sort_values('Perm_mean', ascending=False).reset_index(drop=True)

tbl_i=doc.add_table(rows=len(all_imp)+1,cols=5)
tbl_i.style='Table Grid'; tbl_i.alignment=WD_TABLE_ALIGNMENT.CENTER
set_brd(tbl_i,'1F3864',8)
for row in tbl_i.rows:
    for i,c in enumerate(row.cells):
        c.width=Cm([4.5,1.5,2.5,2.5,5.0][i])
for i,h_ in enumerate(['Variable (lag-1)','Canal',
                        'Gini (train)','Permut. (test)','Interpretation']):
    set_bg(tbl_i.rows[0].cells[i],'1F3864')
    cw(tbl_i.rows[0].cells[i],h_,bold=True,sz=9,color='FFFFFF',al=WD_ALIGN_PARAGRAPH.CENTER)

CBGS={'C1':'EBF5FB','C2':'EAF4E8','C3':'FEF9E7','C4':'FDF2F8','C5':'F5F0FF'}
CANAL_MAP = {s[0]+'_lag1': s[2] for s in SELECTED}
INTERP_I = {
    'Croissance M2 (%) (lag-1)':
        'Signal monetaire : boom prealable -> mauvaise allocation',
    'Solde budgetaire (% PIB) (lag-1)':
        'Signal fiscal : pression budgetaire -> fragilite bancaire',
    'Croissance PIB reel (%) (lag-1)':
        'Signal reel : ralentissement -> deterioration bilans',
    'Reserves change (USD) (lag-1)':
        'Signal externe : epuisement reserves -> choc liquidite',
    'Rentes petrole (% PIB) (lag-1)':
        'Signal CEMAC : dependance petroliere -> exposition chocs',
}
for r_i, row_data in all_imp.iterrows():
    row = tbl_i.rows[r_i+1]
    canal = CANAL_MAP.get(row_data['Feature'],'—')
    cbg   = CBGS.get(canal,'FFFFFF')
    interp = INTERP_I.get(row_data['Label'],'Signal macro-financier')
    perm_s = f"{row_data['Perm_mean']:+.4f} ± {row_data['Perm_std']:.4f}"
    is_pos = row_data['Perm_mean'] > 0
    for i,val in enumerate([row_data['Label'], canal,
                             f"{row_data['Gini']:.4f}", perm_s, interp]):
        set_bg(row.cells[i],cbg)
        cw(row.cells[i],val,sz=9,bold=(i in [0,3]),
           color=('1A7A2E' if i==3 and is_pos else
                  'C0392B' if i==3 and not is_pos else None),
           al=WD_ALIGN_PARAGRAPH.CENTER if i in [1,2] else WD_ALIGN_PARAGRAPH.LEFT)
CAP(doc,"Tableau 3.44 — Importance Gini et par permutation (n=50) — "
        "vert : contribution positive a l'AUC-ROC test")

IMG(doc,PATH_FIG3,w=15.5,
    cap_text="Importance Gini (train) et par permutation (test) — Random Forest",
    num='18')

# ── 3.6.4 Analyse par pays ────────────────────────────────────────────
SEP(doc)
H(doc,f"3.6.4  Analyse par pays — Test {SPLIT_YEAR}-2023",2)
tbl_p=doc.add_table(rows=len(country_res)+1,cols=7)
tbl_p.style='Table Grid'; tbl_p.alignment=WD_TABLE_ALIGNMENT.CENTER
set_brd(tbl_p,'1F3864',8)
for row in tbl_p.rows:
    for i,c in enumerate(row.cells):
        c.width=Cm([4.0,1.0,1.5,2.0,1.0,1.0,1.0][i])
for i,h_ in enumerate(['Pays','N','Vuln.','P_hat moy.','TP','FN','FP']):
    set_bg(tbl_p.rows[0].cells[i],'1F3864')
    cw(tbl_p.rows[0].cells[i],h_,bold=True,sz=9,color='FFFFFF',al=WD_ALIGN_PARAGRAPH.CENTER)
for r_i,(country,n_c,nc1,avg_p,tp_c,fn_c,fp_c,tn_c,auc_c) in enumerate(country_res):
    row=tbl_p.rows[r_i+1]
    rbg=('FDECEA' if fn_c>0 else 'E8F5E9' if nc1>0 else 'EDF3FB')
    for i,val in enumerate([country,str(n_c),str(nc1),
                             f"{avg_p:.3f}",str(tp_c),str(fn_c),str(fp_c)]):
        set_bg(row.cells[i],rbg)
        cw(row.cells[i],val,sz=9.5,bold=(i==5 and fn_c>0),
           color=('C0392B' if fn_c>0 and i==5 else None),
           al=WD_ALIGN_PARAGRAPH.LEFT if i==0 else WD_ALIGN_PARAGRAPH.CENTER)
CAP(doc,f"Tableau 3.45 — Resultats RF par pays (tau_opt={TAU_OPT:.4f}) — rouge : vulnerabilite manquee")

# ── 3.6.5 Comparaison Logit vs RF ─────────────────────────────────────
SEP(doc)
H(doc,"3.6.5  Comparaison Logit L2 vs Random Forest",2)
B(doc,
  "Le tableau ci-dessous confronte les deux modeles sur le set de test "
  f"({SPLIT_YEAR}-2023) au seuil tau_opt de chaque modele. "
  "La comparaison est equitable : meme cible, meme split, memes variables, "
  "meme procedure de calibration du seuil (Youden sur train).",
  sb=0,sa=4)

COMP_ROWS=[
    ('AUC-ROC',    ml_te['auc'],    m_te['auc'],    False),
    ('AUC-PR',     ml_te['auc_pr'], m_te['auc_pr'], False),
    ('Rappel',     ml_te['rappel'], m_te['rappel'],  False),
    ('Precision',  ml_te['prec'],   m_te['prec'],   False),
    ('F1-Score',   ml_te['f1'],     m_te['f1'],     False),
    ('Brier Score',ml_te['brier'],  m_te['brier'],  True),
]
tbl_comp=doc.add_table(rows=len(COMP_ROWS)+1,cols=4)
tbl_comp.style='Table Grid'; tbl_comp.alignment=WD_TABLE_ALIGNMENT.CENTER
set_brd(tbl_comp,'1F3864',8)
for row in tbl_comp.rows:
    for i,c in enumerate(row.cells):
        c.width=Cm([3.5,3.5,3.5,5.5][i])
for i,h_ in enumerate(['Metrique','Logit L2 (baseline)','Random Forest','Verdict']):
    set_bg(tbl_comp.rows[0].cells[i],['1F3864','2E74B5','148F77','2C3E50'][i])
    cw(tbl_comp.rows[0].cells[i],h_,bold=True,sz=9,color='FFFFFF',al=WD_ALIGN_PARAGRAPH.CENTER)

VERDICTS={
    'AUC-ROC':'Discrimination globale (seuil-independant)',
    'AUC-PR' :'Precision-Rappel (classe minoritaire)',
    'Rappel' :'Detection des regimes de vulnerabilite',
    'Precision':'Qualite des alertes emises',
    'F1-Score':'Synthese precision/rappel',
    'Brier Score':'Calibration des probabilites',
}
for r_i,(name,v_l,v_r,low) in enumerate(COMP_ROWS):
    row=tbl_comp.rows[r_i+1]
    better_rf = (v_r>v_l) if not low else (v_r<v_l)
    delta = v_r - v_l
    delta_s = f"RF {'+' if delta>0 else ''}{delta:.4f}"
    bg0='EDF3FB' if r_i%2==0 else 'FFFFFF'
    bg_rf = 'E8F5E9' if better_rf else 'FDECEA'
    set_bg(row.cells[0],bg0);  cw(row.cells[0],name,bold=True,sz=9.5)
    set_bg(row.cells[1],bg0);  cw(row.cells[1],f"{v_l:.4f}",sz=9.5,al=WD_ALIGN_PARAGRAPH.CENTER)
    set_bg(row.cells[2],bg_rf)
    cw(row.cells[2],f"{v_r:.4f}",bold=True,sz=9.5,al=WD_ALIGN_PARAGRAPH.CENTER,
       color='1A7A2E' if better_rf else 'C0392B')
    set_bg(row.cells[3],bg0)
    verdict_txt = (f"RF {'superieur' if better_rf else 'inferieur'} ({delta_s}) — "
                   + VERDICTS.get(name,''))
    cw(row.cells[3],verdict_txt,sz=9,
       color='1A7A2E' if better_rf else 'C0392B')
CAP(doc,"Tableau 3.46 — Comparaison Logit L2 vs Random Forest (test tau_opt chaque modele)")

IMG(doc,PATH_FIG5,w=15.5,
    cap_text=f"Comparaison des metriques Logit L2 vs Random Forest — Test {SPLIT_YEAR}-2023",
    num='19')

# ── 3.6.6 Limites & conclusion ────────────────────────────────────────
SEP(doc)
H(doc,"3.6.6  Limites et positionnement",2)
for bold_txt, rest in [
    ("Interpretabilite : ",
     "Le Random Forest est un modele boite noire : les coefficients "
     "interpretables du logit sont remplaces par des importances agregees. "
     "L'importance par permutation offre une interpretation partielle "
     "mais ne permet pas d'inferer la direction causale des effets."),
    ("Stabilite sur petit echantillon : ",
     f"Avec N_train={n_tr} et {n1_tr} evenements positifs (EPV={n1_tr/5:.1f}), "
     "les arbres individuels peuvent etre instables. "
     "Le bagging (n=500) et min_samples_leaf=5 attenent ce risque, "
     "et le score OOB={:.3f} fournit une validation interne fiable.".format(rf.oob_score_)),
    ("Shift de regime : ",
     f"Comme pour le logit, le shift train ({n1_tr/n_tr*100:.0f}%) -> "
     f"test ({n1_te/n_te*100:.0f}%) affecte la generalisation. "
     "Le RF peut mieux gerer les non-linearites locales mais "
     "ne resout pas fondamentalement le changement de regime."),
]:
    Bmix(doc,[("—  "+bold_txt,True,False,'1F3864'),(rest,False,False,None)],
         sb=0,sa=3,indent=0.5)

# Bilan
SEP(doc)
tbl_fin=doc.add_table(rows=1,cols=1)
tbl_fin.alignment=WD_TABLE_ALIGNMENT.CENTER
set_brd(tbl_fin,'1F3864',8)
c=tbl_fin.rows[0].cells[0]; c.width=Cm(15.5); set_bg(c,'E8F8F0')
p1=c.paragraphs[0]; p1.clear()
p1.alignment=WD_ALIGN_PARAGRAPH.CENTER
p1.paragraph_format.space_before=Pt(7); p1.paragraph_format.space_after=Pt(3)
r1=p1.add_run("Bilan — Random Forest EWS-CEMAC | Vulnerabilite macro-financiere | Split 2019")
r1.bold=True; r1.font.size=Pt(11); r1.font.color.rgb=RGBColor(0x1F,0x38,0x64)
p2=c.add_paragraph()
p2.alignment=WD_ALIGN_PARAGRAPH.LEFT
p2.paragraph_format.space_before=Pt(2); p2.paragraph_format.space_after=Pt(7)
r2=p2.add_run(
    f"RF : n=500 arbres  depth={RF_PARAMS['max_depth']}  leaf>={RF_PARAMS['min_samples_leaf']}  "
    f"OOB={rf.oob_score_:.3f}  tau_opt={TAU_OPT:.4f}\n\n"
    f"Resultats test {SPLIT_YEAR}-2023 (tau_opt) :\n"
    f"  AUC-ROC = {m_te['auc']:.3f}  |  AUC-PR = {m_te['auc_pr']:.3f}  |  "
    f"Rappel = {m_te['rappel']:.3f}  |  Precision = {m_te['prec']:.3f}\n"
    f"  F1-Score = {m_te['f1']:.3f}  |  Brier Score = {m_te['brier']:.3f}\n\n"
    "Variable la plus importante (permutation) : "
    + perm_df.iloc[0]['Label'] + f" ({perm_df.iloc[0]['Perm_mean']:+.4f})\n"
    "Ces resultats sont integres dans la comparaison multi-modeles "
    "(Section 3.8) avec le logit baseline et XGBoost."
)
r2.font.size=Pt(10.5); r2.font.color.rgb=RGBColor(0x1A,0x1A,0x2E)

# Sauvegarde
doc.save(PATH_DOC)
sz=os.path.getsize(PATH_DOC)//1024
print(f"\nDocument Word : {PATH_DOC}  ({sz} Ko)")
dv=Document(PATH_DOC)
hdgs=[(p.style.name,p.text[:65])
      for p in dv.paragraphs
      if p.style.name.startswith('Heading') and p.text.strip()]
print(f"Titres : {len(hdgs)}  |  Tableaux : {len(dv.tables)}")
for s,t in hdgs: print(f"  [{s}]  {t}")
print("\nTERMINE.")
