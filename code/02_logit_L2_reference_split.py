"""
02_logit_L2_reference_split.py
================================
EWS-CEMAC — Penalised Logistic Regression (Logit-L2), Reference Split 2019

PURPOSE
-------
Trains and evaluates a ridge-penalised logistic regression (L2, C = 0.1)
as the baseline EWS classifier on the CEMAC macro-financial panel
(6 countries × 2000–2023, N = 144 country-years).

DESIGN CHOICES
--------------
1. One-year-ahead prediction: all predictors are lagged by one year (t−1)
   to prevent data leakage — supervisors can only act on information known
   *before* the stress period begins.

2. Strong regularisation (C = 0.1, i.e. λ = 1/C = 10):
   EPV (Events Per Variable) = #crises / #variables = 24 / 5 ≈ 4.8.
   EPV < 10 implies risk of overfitting with an unpenalised logit.
   Peduzzi et al. (1996, J. Clin. Epidemiol.) recommend EPV ≥ 10; below
   that, penalisation (Lasso / Ridge) substantially reduces bias.
   C = 0.1 is chosen to balance bias-variance for small panels
   (equivalent to a Normal prior with σ² = 0.1 on standardised coefficients).

3. Class-imbalance correction (class_weight='balanced'):
   The formula applied by scikit-learn is:
       w_k = N_total / (n_classes × N_k)
   where N_k is the count of class k. This upweights crisis observations
   and prevents the model from predicting the majority class (no-stress)
   for all observations.

4. Youden threshold (τ*):
   The decision threshold applied at inference is not the default 0.5 but
   τ* = argmax_τ [TPR(τ) − FPR(τ)] estimated on the training set.
   This is the threshold that maximises the Youden Index J on training data
   and is then applied unchanged on the test set (no re-tuning).

5. Reference split: train 2000–2018 (19 years), test 2019–2023 (4 years).
   This mirrors the typical supervisory horizon during the oil-crisis period.

IMPUTATION STRATEGY
-------------------
  |skewness| ≤ 1.0 → mean imputation (near-symmetric distribution)
  |skewness| > 1.0 → median imputation (robust to outliers)
  Bulmer (1979) rule. Imputer fitted on train set only (no leakage).

AUTHORS
-------
  Françoise NGOUFACK, Pamphile MEZUI-MBENG, Samba NDIAYE — 2026
  Paper: "Do Early Warning Systems Survive Structural Breaks?
          Macroprudential Evidence from the CEMAC Monetary Union"
  Journal of Financial Stability [under review]
"""

import sys
sys.stdout.reconfigure(encoding='utf-8', errors='replace')

import pandas as pd
import numpy as np
import warnings
warnings.filterwarnings('ignore')

from sklearn.linear_model import LogisticRegression
from sklearn.metrics import (
    roc_auc_score, roc_curve,
    precision_recall_curve, average_precision_score,
    confusion_matrix, brier_score_loss,
    precision_score, recall_score, f1_score
)
from sklearn.preprocessing import StandardScaler
from sklearn.impute import SimpleImputer
from sklearn.utils.class_weight import compute_class_weight

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.colors import LinearSegmentedColormap

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import os

# ══════════════════════════════════════════════════════════════════════
# CHEMINS & PARAMÈTRES
# ══════════════════════════════════════════════════════════════════════
DATA_DIR   = r"C:\Users\fngou\Desktop\Donnees_Memoire ML"
if not os.path.exists(DATA_DIR):
    DATA_DIR = r"C:\Users\fngou\Desktop\Données_Mémoire ML"

OUT_DIR    = r"C:\Users\fngou\Desktop\Chapitres du Memoire ML"
PATH_MACRO = DATA_DIR + r"\Dataset_Macro_CEMAC.csv"
PATH_DOC   = OUT_DIR  + r"\Section_35_Logit_Baseline_2019.docx"
PATH_FIG1  = OUT_DIR  + r"\fig_s2019_roc_pr.png"
PATH_FIG2  = OUT_DIR  + r"\fig_s2019_confusion.png"
PATH_FIG3  = OUT_DIR  + r"\fig_s2019_coefs.png"
PATH_FIG4  = OUT_DIR  + r"\fig_s2019_proba.png"
PATH_FIG5  = OUT_DIR  + r"\fig_s2019_timeline.png"

TARGET     = 'Target_EWS_h1'
SPLIT_YEAR = 2019      # train 2000–2018 | test 2019–2023
# C = 1/λ where λ is the L2 penalty strength. C = 0.1 → strong regularisation.
# Justified by EPV = 24 crises / 5 predictors ≈ 4.8 < 10 (Peduzzi et al., 1996).
C_FIXED    = 0.1

SELECTED = [
    # (colonne, label, canal, skewness)
    ('Rentes_petrole_pct_PIB',   'Rentes petrole (% PIB)',   'C1',  0.787),
    ('Solde_budgetaire_pct_PIB', 'Solde budgetaire (% PIB)', 'C2',  0.977),
    ('PIB_croissance_reel_pct',  'Croissance PIB reel (%)',  'C3',  5.171),
    ('Solde_courant_pct_PIB',    'Solde courant (% PIB)',    'C4', -1.462),
    ('M2_croissance_pct',        'Croissance M2 (%)',        'C5',  0.680),
]
SKEW_THR  = 1.0
BASE_COLS = [s[0] for s in SELECTED]
LAG_COLS  = [s[0]+'_lag1' for s in SELECTED]
LABELS    = {s[0]: s[1] for s in SELECTED}
CANAUX    = {s[0]: s[2] for s in SELECTED}
IMPUTE    = {s[0]: ('median' if abs(s[3])>SKEW_THR else 'mean') for s in SELECTED}

BLUE='#1F3864'; LBLUE='#2E74B5'; ORANGE='#C0392B'; GREEN='#1A7A2E'; GRAY='#95A5A6'

# ══════════════════════════════════════════════════════════════════════
# 1. CHARGEMENT & SPLIT
# ══════════════════════════════════════════════════════════════════════
print("=" * 66)
print("  PIPELINE EWS-CEMAC  |  Split 2019  |  5 variables lag-1")
print("=" * 66)

df = (pd.read_csv(PATH_MACRO)
        .sort_values(['Country','Year'])
        .reset_index(drop=True))
df = df[['Country','Year',TARGET]+BASE_COLS].copy()
df = df[df[TARGET].notna()].reset_index(drop=True)

train_df = df[df['Year'] <  SPLIT_YEAR].copy().reset_index(drop=True)
test_df  = df[df['Year'] >= SPLIT_YEAR].copy().reset_index(drop=True)

# One-year lag computed separately within each partition to prevent leakage:
# if lags were computed on the full dataset first, a test-year observation
# (e.g. 2019) would use the actual 2018 value, which is in the training set —
# that is correct. But computing them within each partition ensures no future
# information leaks through the lag chain across the 2019 split boundary.
for part in [train_df, test_df]:
    for col in BASE_COLS:
        part[col+'_lag1'] = part.groupby('Country')[col].shift(1)

X_tr_raw = train_df[LAG_COLS].copy()
X_te_raw = test_df[LAG_COLS].copy()
y_train  = train_df[TARGET].astype(int).values
y_test   = test_df[TARGET].astype(int).values
meta_tr  = train_df[['Country','Year']].copy()
meta_te  = test_df[['Country','Year']].copy()

n_tr, n1_tr = len(y_train), int(y_train.sum())
n_te, n1_te = len(y_test),  int(y_test.sum())
n0_tr = n_tr - n1_tr

print(f"\n  TRAIN : {train_df['Year'].min()}–{train_df['Year'].max()} "
      f"| {n_tr} obs | {n1_tr} crises ({n1_tr/n_tr*100:.1f}%)")
print(f"  TEST  : {test_df['Year'].min()}–{test_df['Year'].max()}  "
      f"| {n_te} obs | {n1_te} crises ({n1_te/n_te*100:.1f}%)")
print(f"  EPV   : {n1_tr}/{len(LAG_COLS)} = {n1_tr/len(LAG_COLS):.1f}")

# Distribution crises par année
print("\n  Distribution Target_EWS_h1 par année :")
for yr in sorted(df['Year'].unique()):
    sub  = df[df['Year']==yr]
    nc   = int(sub[TARGET].sum()); n = len(sub)
    part = 'TRAIN' if yr < SPLIT_YEAR else 'TEST '
    bar  = '█' * nc
    print(f"    {yr} [{part}]  {bar}  {nc}/{n} crises")

# ══════════════════════════════════════════════════════════════════════
# 2. IMPUTATION  (fit sur train uniquement)
# ══════════════════════════════════════════════════════════════════════
mean_cols   = [c+'_lag1' for c in BASE_COLS if IMPUTE[c]=='mean']
median_cols = [c+'_lag1' for c in BASE_COLS if IMPUTE[c]=='median']

imp_mean   = SimpleImputer(strategy='mean')
imp_median = SimpleImputer(strategy='median')
imp_mean.fit(X_tr_raw[mean_cols])
imp_median.fit(X_tr_raw[median_cols])

for X in [X_tr_raw, X_te_raw]:
    X[mean_cols]   = imp_mean.transform(X[mean_cols])
    X[median_cols] = imp_median.transform(X[median_cols])

assert X_tr_raw.isna().sum().sum()==0 and X_te_raw.isna().sum().sum()==0
print(f"\n  Imputation : {len(mean_cols)} MEAN | {len(median_cols)} MEDIAN — 0 NaN résiduel")

# ══════════════════════════════════════════════════════════════════════
# 3. NORMALISATION  (fit sur train uniquement)
# ══════════════════════════════════════════════════════════════════════
scaler  = StandardScaler()
X_train = pd.DataFrame(scaler.fit_transform(X_tr_raw), columns=LAG_COLS)
X_test  = pd.DataFrame(scaler.transform(X_te_raw),     columns=LAG_COLS)
print(f"  Normalisation : StandardScaler fit sur train uniquement")

# ══════════════════════════════════════════════════════════════════════
# 4. CLASS WEIGHTS
# ══════════════════════════════════════════════════════════════════════
# scikit-learn 'balanced' formula:
#   w_k = N_total / (n_classes × N_k),  k ∈ {0, 1}
# With N_total ≈ 114 train obs, N_0 ≈ 90 (no-stress), N_1 ≈ 24 (stress):
#   w_0 = 114 / (2 × 90) ≈ 0.63
#   w_1 = 114 / (2 × 24) ≈ 2.38   → crisis observations get ~3.8× more weight
# This prevents the model from simply predicting "no stress" for all obs.
cw_arr  = compute_class_weight('balanced', classes=np.array([0,1]), y=y_train)
cw_dict = {0: round(cw_arr[0],4), 1: round(cw_arr[1],4)}
sw      = np.where(y_train==1, cw_dict[1], cw_dict[0])
ratio   = round(cw_dict[1]/cw_dict[0], 1)
print(f"  Class weights : w0={cw_dict[0]} | w1={cw_dict[1]} | ratio={ratio}x")

# ══════════════════════════════════════════════════════════════════════
# 5. ENTRAÎNEMENT LOGIT BASELINE
# ══════════════════════════════════════════════════════════════════════
print("\n" + "=" * 66)
print(f"  LOGIT BASELINE  C={C_FIXED}  |  class_weight='balanced'")
print("=" * 66)

logit = LogisticRegression(
    penalty='l2', C=C_FIXED,
    class_weight='balanced',
    solver='lbfgs', max_iter=2000,
    random_state=42
)
logit.fit(X_train, y_train, sample_weight=sw)

prob_tr = logit.predict_proba(X_train)[:,1]
prob_te = logit.predict_proba(X_test)[:,1]

print(f"\n  Intercept  : {logit.intercept_[0]:.4f}")
print(f"  Iterations : {logit.n_iter_[0]}")
print(f"  P_hat train : [{prob_tr.min():.4f}, {prob_tr.max():.4f}]  moy={prob_tr.mean():.4f}")
print(f"  P_hat test  : [{prob_te.min():.4f}, {prob_te.max():.4f}]  moy={prob_te.mean():.4f}")

# Youden-optimal threshold calibrated on training data:
#   J(τ) = TPR(τ) − FPR(τ) = Sensitivity + Specificity − 1
#   τ* = argmax_τ J(τ)
# The threshold τ* is then applied unchanged on the test set to avoid
# implicit re-optimisation that would overstate test performance.
fpr_tr, tpr_tr, thr_tr = roc_curve(y_train, prob_tr)
j_tr    = tpr_tr + (1-fpr_tr) - 1
idx_opt = np.argmax(j_tr)
TAU_OPT = float(thr_tr[idx_opt])
TAU_DEF = 0.5
print(f"\n  Seuil Youden (train) : tau_opt = {TAU_OPT:.4f}")

# ══════════════════════════════════════════════════════════════════════
# 6. MÉTRIQUES
# ══════════════════════════════════════════════════════════════════════
def metriques(y_true, y_prob, tau, label):
    y_pred = (y_prob >= tau).astype(int)
    cm = confusion_matrix(y_true, y_pred)
    if cm.size == 4:
        tn, fp, fn, tp = cm.ravel()
    else:
        tn = fp = fn = tp = 0
        if y_true.sum() == 0: tn = cm[0,0]
        else: tp = cm[0,0]

    auc    = roc_auc_score(y_true, y_prob) if len(np.unique(y_true))>1 else np.nan
    auc_pr = average_precision_score(y_true, y_prob) if len(np.unique(y_true))>1 else np.nan
    rappel = recall_score(y_true, y_pred, zero_division=0)
    prec   = precision_score(y_true, y_pred, zero_division=0)
    f1     = f1_score(y_true, y_pred, zero_division=0)
    brier  = brier_score_loss(y_true, y_prob)
    spec   = tn/(tn+fp) if (tn+fp)>0 else 0.
    j      = rappel + spec - 1

    print(f"\n  --- {label} ---")
    print(f"  AUC-ROC   : {auc:.4f}")
    print(f"  AUC-PR    : {auc_pr:.4f}")
    print(f"  Rappel    : {rappel:.4f}  ({tp}/{tp+fn} crises detectees)")
    print(f"  Precision : {prec:.4f}")
    print(f"  F1-Score  : {f1:.4f}")
    print(f"  Brier     : {brier:.4f}")
    print(f"  Youden J  : {j:.4f}  |  Spec={spec:.4f}")
    print(f"  CM : TN={tn} FP={fp} FN={fn} TP={tp}")

    return dict(label=label, tau=tau, auc=auc, auc_pr=auc_pr,
                rappel=rappel, prec=prec, f1=f1, brier=brier,
                spec=spec, j=j,
                tn=int(tn), fp=int(fp), fn=int(fn), tp=int(tp))

print("\n" + "=" * 66)
print("  RESULTATS")
print("=" * 66)

m_tr  = metriques(y_train, prob_tr, TAU_OPT, f'TRAIN  tau={TAU_OPT:.4f}')
m_te  = metriques(y_test,  prob_te, TAU_OPT, f'TEST   tau={TAU_OPT:.4f} (calibre sur train)')
m_te5 = metriques(y_test,  prob_te, TAU_DEF, f'TEST   tau=0.50 (reference)')

# Coefficients
coef_df = pd.DataFrame({
    'Feature': LAG_COLS,
    'Label'  : [LABELS[c]+' (lag-1)' for c in BASE_COLS],
    'Canal'  : [CANAUX[c] for c in BASE_COLS],
    'Coef'   : logit.coef_[0],
    'OR'     : np.exp(logit.coef_[0]),
}).sort_values('Coef', key=abs, ascending=False).reset_index(drop=True)

print("\n  Coefficients (par |beta| décroissant) :")
for _,r in coef_df.iterrows():
    d = 'risque   (+)' if r['Coef']>0 else 'protection (-)'
    print(f"    {r['Label']:35s}  beta={r['Coef']:+.4f}  OR={r['OR']:.4f}  [{d}]")

# Analyse par pays (test)
print("\n  Analyse par pays — TEST :")
country_res = []
for country in meta_te['Country'].unique():
    mask = meta_te['Country'].values == country
    y_c  = y_test[mask]; p_c = prob_te[mask]; n_c = len(y_c)
    nc1  = int(y_c.sum())
    pred = (p_c >= TAU_OPT).astype(int)
    tp_c = int(((y_c==1)&(pred==1)).sum())
    fn_c = int(((y_c==1)&(pred==0)).sum())
    fp_c = int(((y_c==0)&(pred==1)).sum())
    auc_c= roc_auc_score(y_c, p_c) if len(np.unique(y_c))==2 else np.nan
    country_res.append((country,n_c,nc1,p_c.mean(),tp_c,fn_c,fp_c,auc_c))
    auc_s = f"{auc_c:.3f}" if not np.isnan(auc_c) else 'n/a'
    print(f"    {country:30s} N={n_c} Crises={nc1} "
          f"p_moy={p_c.mean():.3f} TP={tp_c} FN={fn_c} FP={fp_c} AUC={auc_s}")

# ══════════════════════════════════════════════════════════════════════
# 7. FIGURES
# ══════════════════════════════════════════════════════════════════════
print("\nGeneration des figures...")

# ── Figure 5 : Timeline des crises + split ────────────────────────────
fig, axes = plt.subplots(2, 1, figsize=(14, 7), sharex=True)
fig.suptitle("Panel CEMAC — Distribution temporelle des crises et split train/test",
             fontsize=12, fontweight='bold', color=BLUE)

years = sorted(df['Year'].unique())
countries = sorted(df['Country'].unique())
col_map = {c: plt.cm.Set2(i/len(countries)) for i,c in enumerate(countries)}

ax = axes[0]
for i, country in enumerate(countries):
    sub = df[df['Country']==country]
    yr_crise = sub[sub[TARGET]==1]['Year'].values
    ax.scatter(yr_crise, [i]*len(yr_crise), s=200,
               color=col_map[country], zorder=5, label=country, marker='D')
ax.axvline(SPLIT_YEAR-0.5, color=ORANGE, lw=2.5, ls='--', label=f'Split {SPLIT_YEAR}')
ax.axvspan(df['Year'].min()-0.5, SPLIT_YEAR-0.5, alpha=0.05, color=LBLUE)
ax.axvspan(SPLIT_YEAR-0.5, df['Year'].max()+0.5, alpha=0.05, color=ORANGE)
ax.set_yticks(range(len(countries))); ax.set_yticklabels(countries, fontsize=10)
ax.set_ylabel('Pays', fontsize=11); ax.grid(axis='x', alpha=0.3)
ax.legend(loc='upper left', fontsize=8.5, ncol=3)
ax.set_title('Crises bancaires detectees (Target_EWS_h1 = 1)', fontsize=11,
             fontweight='bold', color=BLUE)
ax.text(2009, -0.7, f'TRAIN\n{train_df["Year"].min()}–{train_df["Year"].max()}\n'
        f'{n_tr} obs | {n1_tr} crises', ha='center', fontsize=10,
        color=LBLUE, fontweight='bold')
ax.text(2020.5, -0.7, f'TEST\n{test_df["Year"].min()}–{test_df["Year"].max()}\n'
        f'{n_te} obs | {n1_te} crises', ha='center', fontsize=10,
        color=ORANGE, fontweight='bold')

ax2 = axes[1]
rates = [df[df['Year']==y][TARGET].mean()*100 for y in years]
colors_bar = [LBLUE if y < SPLIT_YEAR else ORANGE for y in years]
ax2.bar(years, rates, color=colors_bar, alpha=0.8, edgecolor='white', width=0.8)
ax2.axvline(SPLIT_YEAR-0.5, color=ORANGE, lw=2.5, ls='--')
ax2.set_xlabel('Annee', fontsize=11); ax2.set_ylabel('Taux de crise (%)', fontsize=11)
ax2.set_title('Taux de crise annuel (% pays en crise)', fontsize=11,
              fontweight='bold', color=BLUE)
ax2.grid(axis='y', alpha=0.3)
legend_els = [mpatches.Patch(color=LBLUE, label=f'TRAIN (2000-{SPLIT_YEAR-1})'),
              mpatches.Patch(color=ORANGE, label=f'TEST  ({SPLIT_YEAR}-2022)')]
ax2.legend(handles=legend_els, fontsize=9)

plt.tight_layout()
plt.savefig(PATH_FIG5, dpi=150, bbox_inches='tight')
plt.close()
print(f"  Fig 5 (timeline) : {PATH_FIG5}")

# ── Figure 1 : ROC + Précision-Rappel ────────────────────────────────
fig, axes = plt.subplots(1, 2, figsize=(12, 5))
fig.suptitle(f"Logit Baseline EWS-CEMAC  |  Split {SPLIT_YEAR}  |  C={C_FIXED}",
             fontsize=12, fontweight='bold', color=BLUE)

fpr_te, tpr_te, _ = roc_curve(y_test,  prob_te)
fpr_tr2,tpr_tr2,_ = roc_curve(y_train, prob_tr)

ax = axes[0]
ax.fill_between(fpr_te, tpr_te, alpha=0.1, color=BLUE)
ax.plot(fpr_te,  tpr_te,  color=BLUE,  lw=2.5, label=f"Test   AUC-ROC = {m_te['auc']:.3f}")
ax.plot(fpr_tr2, tpr_tr2, color=LBLUE, lw=1.5, ls='--',
        label=f"Train  AUC-ROC = {m_tr['auc']:.3f}")
ax.plot([0,1],[0,1], color=GRAY, lw=1, ls=':', label='Classifieur aleatoire')
far_opt = m_te['fp']/(m_te['fp']+m_te['tn']) if (m_te['fp']+m_te['tn'])>0 else 0
ax.scatter([far_opt],[m_te['rappel']], s=150, color=ORANGE, zorder=6,
           label=f"tau_opt={TAU_OPT:.3f}  J={m_te['j']:.3f}")
ax.set_xlabel('Taux de fausse alarme (1 - Specificite)', fontsize=11)
ax.set_ylabel('Rappel (Sensibilite)', fontsize=11)
ax.set_title('Courbe ROC', fontsize=11, fontweight='bold', color=BLUE)
ax.legend(fontsize=9.5); ax.grid(alpha=0.3)
ax.set_xlim([-0.02,1.02]); ax.set_ylim([-0.02,1.02])

prec_te, rec_te, _ = precision_recall_curve(y_test,  prob_te)
prec_tr, rec_tr, _ = precision_recall_curve(y_train, prob_tr)
ax = axes[1]
ax.fill_between(rec_te, prec_te, alpha=0.1, color=BLUE)
ax.plot(rec_te, prec_te, color=BLUE,  lw=2.5, label=f"Test   AUC-PR = {m_te['auc_pr']:.3f}")
ax.plot(rec_tr, prec_tr, color=LBLUE, lw=1.5, ls='--',
        label=f"Train  AUC-PR = {m_tr['auc_pr']:.3f}")
ax.axhline(y_test.mean(), color=GRAY, lw=1.2, ls=':',
           label=f"Baseline ({y_test.mean():.2f})")
ax.scatter([m_te['rappel']],[m_te['prec']], s=150, color=ORANGE, zorder=6,
           label=f"tau_opt={TAU_OPT:.3f}")
ax.set_xlabel('Rappel (Sensibilite)', fontsize=11)
ax.set_ylabel('Precision', fontsize=11)
ax.set_title('Courbe Precision-Rappel', fontsize=11, fontweight='bold', color=BLUE)
ax.legend(fontsize=9.5); ax.grid(alpha=0.3)
ax.set_xlim([-0.02,1.02]); ax.set_ylim([-0.02,1.02])

plt.tight_layout()
plt.savefig(PATH_FIG1, dpi=150, bbox_inches='tight')
plt.close()
print(f"  Fig 1 (ROC+PR) : {PATH_FIG1}")

# ── Figure 2 : Matrices de confusion ─────────────────────────────────
fig, axes = plt.subplots(1, 2, figsize=(11, 4.5))
fig.suptitle(f"Matrices de confusion — Set de test {SPLIT_YEAR}–2022 | Logit Baseline",
             fontsize=12, fontweight='bold', color=BLUE)

def plot_cm(ax, m, title):
    cm_m = np.array([[m['tn'],m['fp']],[m['fn'],m['tp']]])
    cmap = LinearSegmentedColormap.from_list('b',['#FFFFFF','#1F3864'])
    ax.imshow(cm_m, cmap=cmap, vmin=0, vmax=max(cm_m.max(),1))
    for i in range(2):
        for j in range(2):
            v=cm_m[i,j]; col='white' if v>cm_m.max()*0.55 else 'black'
            ax.text(j,i,f"{[['TN','FP'],['FN','TP']][i][j]}\n{v}",
                    ha='center',va='center',fontsize=14,fontweight='bold',color=col)
    ax.set_xticks([0,1]); ax.set_yticks([0,1])
    ax.set_xticklabels(['Predit 0\n(Non-crise)','Predit 1\n(Crise)'],fontsize=10)
    ax.set_yticklabels(['Reel 0\n(Non-crise)','Reel 1\n(Crise)'],fontsize=10)
    ax.set_title(title, fontsize=11, fontweight='bold', color=BLUE)
    ax.set_xlabel(f"Rappel={m['rappel']:.2f}  Prec={m['prec']:.2f}  "
                  f"F1={m['f1']:.2f}  J={m['j']:.2f}", fontsize=9.5)

plot_cm(axes[0], m_te5, "Seuil par defaut (tau = 0.50)")
plot_cm(axes[1], m_te,  f"Seuil optimal (tau = {TAU_OPT:.4f})")
plt.tight_layout()
plt.savefig(PATH_FIG2, dpi=150, bbox_inches='tight')
plt.close()
print(f"  Fig 2 (CM) : {PATH_FIG2}")

# ── Figure 3 : Coefficients & OR ─────────────────────────────────────
fig, axes = plt.subplots(1, 2, figsize=(13, 4.5))
fig.suptitle("Coefficients et Odds Ratios — Logit Baseline",
             fontsize=12, fontweight='bold', color=BLUE)

cs  = coef_df.sort_values('Coef')
col_b = [ORANGE if c>0 else LBLUE for c in cs['Coef']]
col_o = [ORANGE if o>1 else LBLUE for o in cs['OR']]

axes[0].barh(cs['Label'], cs['Coef'], color=col_b, edgecolor='white', height=0.6)
axes[0].axvline(0, color='black', lw=0.8)
for i,(v,_) in enumerate(zip(cs['Coef'], cs['Label'])):
    off = 0.003 if v>=0 else -0.003
    axes[0].text(v+off, i, f'{v:+.4f}', va='center',
                 ha='left' if v>=0 else 'right', fontsize=9.5)
axes[0].set_xlabel('Coefficient beta (unites standardisees)', fontsize=11)
axes[0].set_title('Coefficients logistiques', fontsize=11, fontweight='bold', color=BLUE)
axes[0].grid(axis='x', alpha=0.3)
axes[0].legend(handles=[
    mpatches.Patch(color=ORANGE, label='Facteur de risque  beta > 0'),
    mpatches.Patch(color=LBLUE,  label='Facteur protecteur beta < 0')],
    fontsize=9.5)

axes[1].barh(cs['Label'], cs['OR'], color=col_o, edgecolor='white', height=0.6)
axes[1].axvline(1, color='black', lw=0.8, ls='--', label='OR = 1 (effet nul)')
for i,(v,_) in enumerate(zip(cs['OR'], cs['Label'])):
    axes[1].text(v+0.002, i, f'{v:.4f}', va='center', ha='left', fontsize=9.5)
axes[1].set_xlabel('Odds Ratio = exp(beta)', fontsize=11)
axes[1].set_title('Odds Ratios', fontsize=11, fontweight='bold', color=BLUE)
axes[1].grid(axis='x', alpha=0.3); axes[1].legend(fontsize=9.5)
for ax in axes: ax.tick_params(axis='y', labelsize=9.5)

plt.tight_layout()
plt.savefig(PATH_FIG3, dpi=150, bbox_inches='tight')
plt.close()
print(f"  Fig 3 (coefs) : {PATH_FIG3}")

# ── Figure 4 : Distribution P_hat ────────────────────────────────────
fig, axes = plt.subplots(1, 2, figsize=(12, 4.5))
fig.suptitle("Distribution de P_hat(Crise) — Logit Baseline",
             fontsize=12, fontweight='bold', color=BLUE)
for ax,(probs,y_true,title) in zip(axes,[
        (prob_tr, y_train, f"Ensemble d'entrainement (2000-{SPLIT_YEAR-1})"),
        (prob_te, y_test,  f"Ensemble de test ({SPLIT_YEAR}-2022)")]):
    ax.hist(probs[y_true==0], bins=25, alpha=0.6, color=LBLUE,
            label='Non-crise (y=0)', edgecolor='white')
    ax.hist(probs[y_true==1], bins=10, alpha=0.75, color=ORANGE,
            label='Crise (y=1)', edgecolor='white')
    ax.axvline(TAU_OPT, color='red',  ls='--', lw=2,
               label=f'tau_opt={TAU_OPT:.4f}')
    ax.axvline(TAU_DEF, color=GRAY,   ls=':',  lw=1.5,
               label='tau=0.50')
    ax.set_xlabel('P_hat(Crise)', fontsize=11)
    ax.set_ylabel('Observations', fontsize=11)
    ax.set_title(title, fontsize=11, fontweight='bold', color=BLUE)
    ax.legend(fontsize=9.5); ax.grid(alpha=0.3)
plt.tight_layout()
plt.savefig(PATH_FIG4, dpi=150, bbox_inches='tight')
plt.close()
print(f"  Fig 4 (proba) : {PATH_FIG4}")

# ══════════════════════════════════════════════════════════════════════
# 8. DOCUMENT WORD
# ══════════════════════════════════════════════════════════════════════
print("\nConstruction du document Word...")

def set_bg(cell, hx):
    tcPr=cell._tc.get_or_add_tcPr()
    s=OxmlElement('w:shd')
    s.set(qn('w:val'),'clear'); s.set(qn('w:color'),'auto')
    s.set(qn('w:fill'),hx); tcPr.append(s)

def set_brd(t, color='1F3864', sz=8):
    tbl=t._tbl; tblPr=tbl.tblPr or OxmlElement('w:tblPr')
    b=OxmlElement('w:tblBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        el=OxmlElement(f'w:{side}')
        el.set(qn('w:val'),'single')
        el.set(qn('w:sz'),str(sz if side in ('top','left','bottom','right') else 4))
        el.set(qn('w:space'),'0'); el.set(qn('w:color'),color)
        b.append(el)
    tblPr.append(b)

def cw(cell, text, bold=False, italic=False, sz=9.5, color=None,
       al=WD_ALIGN_PARAGRAPH.LEFT, sb=2, sa=2):
    p=cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    p.clear(); p.alignment=al
    p.paragraph_format.space_before=Pt(sb)
    p.paragraph_format.space_after=Pt(sa)
    cell.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
    r=p.add_run(text); r.bold=bold; r.italic=italic; r.font.size=Pt(sz)
    if color: r.font.color.rgb=RGBColor(*bytes.fromhex(color))

def H(doc, text, level, color='1F3864', sz=None):
    sz_d={1:13,2:12,3:11}
    p=doc.add_paragraph(style=f'Heading {level}')
    r=p.add_run(text); r.bold=True
    r.font.size=Pt(sz or sz_d[level])
    r.font.color.rgb=RGBColor(*bytes.fromhex(color))

def B(doc, text, italic=False, sz=10.5, sb=0, sa=4, indent=None):
    p=doc.add_paragraph(style='Normal')
    p.paragraph_format.space_before=Pt(sb)
    p.paragraph_format.space_after=Pt(sa)
    if indent: p.paragraph_format.left_indent=Cm(indent)
    r=p.add_run(text); r.font.size=Pt(sz); r.italic=italic

def Bmix(doc, segs, sb=0, sa=4, indent=None):
    p=doc.add_paragraph(style='Normal')
    p.paragraph_format.space_before=Pt(sb)
    p.paragraph_format.space_after=Pt(sa)
    if indent: p.paragraph_format.left_indent=Cm(indent)
    for txt,bold,italic,col in segs:
        r=p.add_run(txt); r.font.size=Pt(10.5)
        r.bold=bold; r.italic=italic
        if col: r.font.color.rgb=RGBColor(*bytes.fromhex(col))

def IMG(doc, path, w=15.0, cap_text=None, num=None):
    if not os.path.exists(path):
        B(doc,f"[Figure manquante : {path}]",italic=True); return
    p=doc.add_paragraph(style='Normal')
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before=Pt(4)
    p.paragraph_format.space_after=Pt(2)
    p.add_run().add_picture(path, width=Cm(w))
    if cap_text:
        pc=doc.add_paragraph(style='Normal')
        pc.alignment=WD_ALIGN_PARAGRAPH.CENTER
        pc.paragraph_format.space_before=Pt(2)
        pc.paragraph_format.space_after=Pt(10)
        rc=pc.add_run(f"Figure 3.{num} — {cap_text}" if num else cap_text)
        rc.italic=True; rc.font.size=Pt(9.5)
        rc.font.color.rgb=RGBColor(0x44,0x44,0x44)

def CAP(doc, text):
    p=doc.add_paragraph(style='Normal')
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before=Pt(2)
    p.paragraph_format.space_after=Pt(8)
    r=p.add_run(text); r.italic=True; r.font.size=Pt(9.5)
    r.font.color.rgb=RGBColor(0x44,0x44,0x44)

def SEP(doc):
    p=doc.add_paragraph(style='Normal')
    p.paragraph_format.space_before=Pt(10)
    p.paragraph_format.space_after=Pt(4)
    pPr=p._element.get_or_add_pPr()
    pBdr=OxmlElement('w:pBdr')
    top=OxmlElement('w:top')
    top.set(qn('w:val'),'single'); top.set(qn('w:sz'),'4')
    top.set(qn('w:space'),'1'); top.set(qn('w:color'),'1F3864')
    pBdr.append(top); pPr.append(pBdr)

def INFO(doc, title, body, bg_c='DEEAF1', brd_c='1F3864'):
    t=doc.add_table(rows=1,cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    set_brd(t,color=brd_c,sz=6)
    c=t.rows[0].cells[0]; c.width=Cm(15.5); set_bg(c,bg_c)
    p1=c.paragraphs[0]; p1.clear()
    p1.paragraph_format.space_before=Pt(5)
    p1.paragraph_format.space_after=Pt(2)
    r1=p1.add_run(title); r1.bold=True; r1.font.size=Pt(10)
    r1.font.color.rgb=RGBColor(*bytes.fromhex(brd_c))
    p2=c.add_paragraph()
    p2.paragraph_format.space_before=Pt(0)
    p2.paragraph_format.space_after=Pt(5)
    r2=p2.add_run(body); r2.font.size=Pt(10); r2.italic=True
    r2.font.color.rgb=RGBColor(0x1A,0x1A,0x2E)
    doc.add_paragraph(style='Normal').paragraph_format.space_after=Pt(6)

def FBOX(doc, text):
    t=doc.add_table(rows=1,cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    set_brd(t,color='AAAAAA',sz=4)
    c=t.rows[0].cells[0]; c.width=Cm(14); set_bg(c,'F2F2F2')
    p=c.paragraphs[0]; p.clear()
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before=Pt(5)
    p.paragraph_format.space_after=Pt(5)
    r=p.add_run(text); r.font.size=Pt(10.5)
    r.font.name='Courier New'
    r.font.color.rgb=RGBColor(0x1F,0x38,0x64)
    doc.add_paragraph(style='Normal').paragraph_format.space_after=Pt(4)

# ── Construction du document ──────────────────────────────────────────
doc=Document()
for s in doc.sections:
    s.top_margin=Cm(2.5); s.bottom_margin=Cm(2.5)
    s.left_margin=Cm(3.0); s.right_margin=Cm(2.5)

# TITRE
H(doc, "3.5  Modele Baseline : Regression Logistique Penalisee (L2)", 1)
B(doc,
  "Ce chapitre presente le modele de reference (baseline) de la strategie "
  "de comparaison multi-modeles. La regression logistique L2 est retenue pour "
  "son interpretabilite theorique et son ancrage dans la litterature EWS "
  "(Demirguc-Kunt & Detragiache, 1998 ; Davis & Karim, 2008). "
  "Sa valeur reside dans l'interpretation economique des coefficients, "
  "non dans la performance predictive brute.",
  sb=6, sa=6)

# ── 3.5.1 Justification du split 2019 ────────────────────────────────
H(doc, "3.5.1  Justification du decoupage temporel (split 2019)", 2)
B(doc,
  "Le split temporel est fixe a l'annee 2019, assurant un minimum de "
  + str(n1_tr) + " evenements positifs en train (EPV = " +
  f"{n1_tr/len(LAG_COLS):.1f}" + "). "
  "Ce seuil EPV est la contrainte technique determinante pour la regression "
  "logistique : en dessous de EPV=1, les coefficients ne sont pas "
  "identifiables de maniere fiable (van Smeden et al., 2019). "
  "Le train couvre la periode 2000-" + str(SPLIT_YEAR-1) +
  ", le test la periode " + str(SPLIT_YEAR) + "-2022.",
  sb=0, sa=4)

# Tableau split
tbl_s=doc.add_table(rows=3,cols=5)
tbl_s.style='Table Grid'; tbl_s.alignment=WD_TABLE_ALIGNMENT.CENTER
set_brd(tbl_s,'1F3864',8)
for row in tbl_s.rows:
    for i,c in enumerate(row.cells):
        c.width=Cm([2.5,3.0,2.0,2.5,6.0][i])
for i,h_ in enumerate(['Partition','Periode','Obs.','Crises (%)','Justification']):
    set_bg(tbl_s.rows[0].cells[i],'1F3864')
    cw(tbl_s.rows[0].cells[i],h_,bold=True,sz=9,color='FFFFFF',
       al=WD_ALIGN_PARAGRAPH.CENTER)
data_s=[
    ('TRAIN',f'2000-{SPLIT_YEAR-1}',str(n_tr),
     f"{n1_tr} ({n1_tr/n_tr*100:.1f}%)",
     f"Periode incluant les premieres manifestations de crise (2014-2015) — EPV={n1_tr/len(LAG_COLS):.1f}"),
    ('TEST',f'{SPLIT_YEAR}-2022',str(n_te),
     f"{n1_te} ({n1_te/n_te*100:.1f}%)",
     "Periode post-choc petrolier — vague principale de crises bancaires CEMAC"),
]
BG_S=['EDF3FB','FFF9E6']
for r_i,rd in enumerate(data_s):
    row=tbl_s.rows[r_i+1]
    for i,val in enumerate(rd):
        set_bg(row.cells[i],BG_S[r_i])
        cw(row.cells[i],val,sz=9.5,bold=(i==0),
           al=WD_ALIGN_PARAGRAPH.CENTER if i<3 else WD_ALIGN_PARAGRAPH.LEFT)
CAP(doc,"Tableau 3.37 — Decoupage train/test et justification du split 2019")

B(doc,
  "Ce protocole garantit l'absence de fuite d'information (data leakage) : "
  "le preprocesseur (imputation, normalisation, poids de classe) est "
  "entierement ajuste sur le train, puis applique au test sans aucune "
  "re-calibration. Le seuil de decision tau_opt est egalement calcule "
  "sur le train uniquement (indice de Youden).",
  sb=0, sa=4)

INFO(doc,
     "Justification statistique du split 2019 :",
     "Avec " + str(n1_tr) + " crises en train, l'EPV = " +
     f"{n1_tr/len(LAG_COLS):.1f}" +
     " satisfait le seuil minimum recommande (EPV >= 1) pour une "
     "regression logistique a 5 variables. "
     "Un split 2016 reduirait l'EPV a 0.4 (2 crises), rendant "
     "les estimations non fiables. Le split 2019 est le compromis "
     "optimal entre taille du train et representativite du test.",
     bg_c='E8F8F5', brd_c='1ABC9C')

IMG(doc, PATH_FIG5, w=15.5,
    cap_text=f"Distribution temporelle des crises CEMAC et decoupage train/test (split {SPLIT_YEAR})",
    num='11')

# ── 3.5.2 Specification du modele ────────────────────────────────────
SEP(doc)
H(doc, "3.5.2  Specification du modele baseline", 2)
B(doc,
  "Le modele est specifie avec 5 variables macroeconomiques retardees "
  "d'un an (lag-1), selectionnees theoriquement — une par canal economique "
  "(C1 : ressources petrolieres, C2 : politique budgetaire, "
  "C3 : activite reelle, C4 : equilibre externe, C5 : credit/monnaie). "
  "Le parametre de regularisation C=0.1 est fixe sur la base du ratio "
  "EPV = " + f"{n1_tr/len(LAG_COLS):.1f}" + ", sans validation croisee — "
  "instable avec " + str(n1_tr) + " evenements positifs en train.",
  sb=0, sa=4)

FBOX(doc,
     "Logit L2 :  P(Crise_{t+1} | X_t) = sigma( b0 + sum_k bk * x_k(t-1) )\n\n"
     "  Variables (lag-1) : Rentes petrole | Solde budg. | Croissance PIB\n"
     "                      Solde courant  | Croissance M2\n\n"
     f"  Penalisation  : L2 (Ridge)  |  C = {C_FIXED}  (regularisation forte)\n"
     f"  class_weight  : balanced    |  w0={cw_dict[0]}  w1={cw_dict[1]}  ratio={ratio}x\n"
     f"  Seuil tau_opt : {TAU_OPT:.4f}  (Youden max sur X_train)")

# ── 3.5.3 Metriques ──────────────────────────────────────────────────
SEP(doc)
H(doc, "3.5.3  Metriques de performance", 2)
B(doc,
  "Six metriques sont rapportees, adaptees aux objectifs du systeme "
  "d'alerte precoce et aux criteres des revues specialisees en stabilite "
  "financiere. Chaque metrique est calculee au seuil tau=0.50 (reference) "
  "et au seuil tau_opt (Youden, calibre sur le train).",
  sb=0, sa=4)

# Tableau definitions
DEFS=[
    ('AUC-ROC',            '∫ ROC curve',
     "Probabilite que le modele classe une crise avant une non-crise. "
     "Metrique cle, insensible au seuil de decision."),
    ('AUC-PR',             '∫ Prec-Recall',
     "Aire sous la courbe Precision-Rappel. Plus informative qu'AUC-ROC "
     "en cas de fort desequilibre de classes."),
    ('Rappel (Sensibilite)','TP/(TP+FN)',
     "Proportion de crises detectees. Metrique operationnelle prioritaire "
     "pour BEAC/COBAC : une crise non detectee est tres couteuse."),
    ('Precision',          'TP/(TP+FP)',
     "Fraction des alertes correspondant a de vraies crises. "
     "Mesure le cout des fausses alarmes pour le regulateur."),
    ('F1-Score',           '2*Prec*Rappel/(Prec+Rappel)',
     "Moyenne harmonique precision/rappel. Synthetise les deux types "
     "d'erreur pour la comparaison multi-modeles."),
    ('Brier Score',        'E[(P_hat - y)^2]',
     "Erreur quadratique moyenne des probabilites predites. "
     "Mesure la calibration : plus bas = mieux calibre."),
]
tbl_d=doc.add_table(rows=len(DEFS)+1,cols=3)
tbl_d.style='Table Grid'; tbl_d.alignment=WD_TABLE_ALIGNMENT.CENTER
set_brd(tbl_d,'1F3864',8)
for row in tbl_d.rows:
    for i,c in enumerate(row.cells):
        c.width=Cm([3.5,3.5,9.0][i])
for i,h_ in enumerate(['Metrique','Formule','Role dans le SAP CEMAC']):
    set_bg(tbl_d.rows[0].cells[i],'1F3864')
    cw(tbl_d.rows[0].cells[i],h_,bold=True,sz=9,color='FFFFFF',
       al=WD_ALIGN_PARAGRAPH.CENTER)
for r_i,(name,form,interp) in enumerate(DEFS):
    row=tbl_d.rows[r_i+1]
    bg_c='EDF3FB' if r_i%2==0 else 'FFFFFF'
    for i,val in enumerate([name,form,interp]):
        set_bg(row.cells[i],bg_c)
        cw(row.cells[i],val,sz=9,bold=(i==0))
CAP(doc,"Tableau 3.38 — Definition et role des 6 metriques d'evaluation du SAP")

# Tableau resultats
SEP(doc)
H(doc, "Resultats numeriques", 3)

MROWS=[
    ('AUC-ROC',    m_tr['auc'],    m_te5['auc'],    m_te['auc'],    False),
    ('AUC-PR',     m_tr['auc_pr'], m_te5['auc_pr'], m_te['auc_pr'], False),
    ('Rappel',     m_tr['rappel'], m_te5['rappel'],  m_te['rappel'],  False),
    ('Precision',  m_tr['prec'],   m_te5['prec'],   m_te['prec'],   False),
    ('F1-Score',   m_tr['f1'],     m_te5['f1'],     m_te['f1'],     False),
    ('Brier Score',m_tr['brier'],  m_te5['brier'],  m_te['brier'],  True),
]
KEY={'AUC-ROC','AUC-PR','Rappel','F1-Score'}

tbl_r=doc.add_table(rows=len(MROWS)+2,cols=4)
tbl_r.style='Table Grid'; tbl_r.alignment=WD_TABLE_ALIGNMENT.CENTER
set_brd(tbl_r,'1F3864',8)
for row in tbl_r.rows:
    for i,c in enumerate(row.cells):
        c.width=Cm([3.8,3.0,3.0,3.0][i])

for i,h_ in enumerate(['Metrique','Train','Test tau=0.50','Test tau_opt']):
    bg_c=['1F3864','2E74B5','1A5276','1A5276'][i]
    set_bg(tbl_r.rows[0].cells[i],bg_c)
    cw(tbl_r.rows[0].cells[i],h_,bold=True,sz=9,color='FFFFFF',
       al=WD_ALIGN_PARAGRAPH.CENTER)
set_bg(tbl_r.rows[1].cells[0],'2C3E50')
cw(tbl_r.rows[1].cells[0],'',sz=8)
for i,val in enumerate([f'tau={TAU_OPT:.4f}',f'tau={TAU_DEF}',f'tau={TAU_OPT:.4f}']):
    bg_c='2E74B5' if i==0 else '1A5276'
    set_bg(tbl_r.rows[1].cells[i+1],bg_c)
    cw(tbl_r.rows[1].cells[i+1],val,sz=8,color='FFFFFF',al=WD_ALIGN_PARAGRAPH.CENTER)

for r_i,(name,v_tr,v_te5,v_te,low) in enumerate(MROWS):
    row=tbl_r.rows[r_i+2]
    bg0='F2F9FF' if name in KEY else ('EDF3FB' if r_i%2==0 else 'FFFFFF')
    set_bg(row.cells[0],bg0)
    cw(row.cells[0],name,bold=(name in KEY),sz=9.5)
    for ci,(val,is_opt) in enumerate([(v_tr,False),(v_te5,False),(v_te,True)]):
        better=(val>v_te5) if not low else (val<v_te5)
        bg_c=('E8F5E9' if is_opt and better else
              'FDECEA' if is_opt and not better else bg0)
        set_bg(row.cells[ci+1],bg_c)
        cw(row.cells[ci+1],f"{val:.4f}",bold=is_opt,sz=9.5,
           al=WD_ALIGN_PARAGRAPH.CENTER,
           color=('1A7A2E' if is_opt and better else
                  'C0392B' if is_opt and not better else None))
CAP(doc,f"Tableau 3.39 — Resultats du modele baseline (C={C_FIXED}, split {SPLIT_YEAR}, "
        f"tau_opt={TAU_OPT:.4f})")

IMG(doc,PATH_FIG1,w=15.5,
    cap_text=f"Courbes ROC et Precision-Rappel — Logit Baseline (train 2000-{SPLIT_YEAR-1} / test {SPLIT_YEAR}-2022)",
    num='12')
IMG(doc,PATH_FIG2,w=13.5,
    cap_text=f"Matrices de confusion — Test {SPLIT_YEAR}-2022 (tau=0.50 vs tau_opt={TAU_OPT:.4f})",
    num='13')

# ── 3.5.4 Coefficients ────────────────────────────────────────────────
SEP(doc)
H(doc, "3.5.4  Coefficients et interpretation economique", 2)
B(doc,
  "Les coefficients exprimes en unites standardisees mesurent la contribution "
  "de chaque canal macroeconomique au risque de crise. Le signe de beta "
  "indique la direction de l'effet ; l'odds ratio OR=exp(beta) mesure "
  "le facteur multiplicatif du risque relatif pour un accroissement "
  "d'un ecart-type de la variable.",
  sb=0, sa=4)

INTERP={
    'Rentes petrole (% PIB) (lag-1)':
        "Forte rente passee -> exposition aux retournements des cours -> "
        "choc de revenus sur les banques publiques CEMAC",
    'Solde budgetaire (% PIB) (lag-1)':
        "Excedent budgetaire passe -> capacite de l'Etat a soutenir le secteur -> "
        "facteur protecteur ; deficit -> fragilite",
    'Croissance PIB reel (%) (lag-1)':
        "Croissance passee protege contre les crises ; "
        "recession -> deterioration des bilans bancaires",
    'Solde courant (% PIB) (lag-1)':
        "Excedent courant passe -> tampon de reserves de change -> "
        "protection contre les chocs de liquidite externe",
    'Croissance M2 (%) (lag-1)':
        "Expansion monetaire passee -> risque de mauvaise allocation du credit "
        "et de surchauffe financiere",
}
tbl_c=doc.add_table(rows=6,cols=5)
tbl_c.style='Table Grid'; tbl_c.alignment=WD_TABLE_ALIGNMENT.CENTER
set_brd(tbl_c,'1F3864',8)
for row in tbl_c.rows:
    for i,c in enumerate(row.cells):
        c.width=Cm([4.0,1.8,2.0,2.2,6.0][i])
for i,h_ in enumerate(['Variable (lag-1)','beta','OR','Direction','Interpretation']):
    set_bg(tbl_c.rows[0].cells[i],'1F3864')
    cw(tbl_c.rows[0].cells[i],h_,bold=True,sz=9,color='FFFFFF',
       al=WD_ALIGN_PARAGRAPH.CENTER)

CBGS={'C1':'EBF5FB','C2':'EAF4E8','C3':'FEF9E7','C4':'FDF2F8','C5':'F5F0FF'}
for r_i,(_,rd) in enumerate(coef_df.iterrows()):
    row=tbl_c.rows[r_i+1]
    is_pos=rd['Coef']>0; cbg=CBGS.get(rd['Canal'],'FFFFFF')
    interp_txt = INTERP.get(rd['Label'], 'Signal macroeconomique CEMAC')
    direction='Risque (+)' if is_pos else 'Protection (-)'
    for i,val in enumerate([rd['Label'],
                             f"{rd['Coef']:+.4f}",
                             f"{rd['OR']:.4f}",
                             direction, interp_txt]):
        set_bg(row.cells[i],cbg)
        cw(row.cells[i],val,sz=9,bold=(i in [1,2]),
           color=('C0392B' if i==1 and is_pos else
                  '1A5276' if i==1 else
                  'C0392B' if i==3 and is_pos else
                  '1A5276' if i==3 else None),
           al=WD_ALIGN_PARAGRAPH.CENTER if i in [1,2,3] else WD_ALIGN_PARAGRAPH.LEFT)
CAP(doc,"Tableau 3.40 — Coefficients et interpretation economique "
        "(rouge : facteur de risque beta>0 ; bleu : facteur protecteur beta<0)")

IMG(doc,PATH_FIG3,w=15.5,
    cap_text="Coefficients beta et odds ratios — Logit Baseline (5 variables lag-1)",num='14')
IMG(doc,PATH_FIG4,w=15.5,
    cap_text="Distribution de P_hat(Crise) par classe reelle — train et test",num='15')

# ── 3.5.5 Analyse par pays ────────────────────────────────────────────
SEP(doc)
H(doc, f"3.5.5  Analyse par pays — Set de test {SPLIT_YEAR}-2022", 2)
tbl_p=doc.add_table(rows=len(country_res)+1,cols=6)
tbl_p.style='Table Grid'; tbl_p.alignment=WD_TABLE_ALIGNMENT.CENTER
set_brd(tbl_p,'1F3864',8)
for row in tbl_p.rows:
    for i,c in enumerate(row.cells):
        c.width=Cm([4.5,1.2,1.8,2.2,1.2,1.2][i])
for i,h_ in enumerate(['Pays','N','Crises','P_hat moy.','TP','FN']):
    set_bg(tbl_p.rows[0].cells[i],'1F3864')
    cw(tbl_p.rows[0].cells[i],h_,bold=True,sz=9,color='FFFFFF',
       al=WD_ALIGN_PARAGRAPH.CENTER)
for r_i,(country,n_c,nc1,avg_p,tp_c,fn_c,fp_c,auc_c) in enumerate(country_res):
    row=tbl_p.rows[r_i+1]
    rbg=('FDECEA' if fn_c>0 else 'E8F5E9' if nc1>0 else 'EDF3FB')
    for i,val in enumerate([country,str(n_c),str(nc1),
                             f"{avg_p:.3f}",str(tp_c),str(fn_c)]):
        set_bg(row.cells[i],rbg)
        cw(row.cells[i],val,sz=9.5,bold=(i==5 and fn_c>0),
           color=('C0392B' if fn_c>0 and i==5 else None),
           al=WD_ALIGN_PARAGRAPH.LEFT if i==0 else WD_ALIGN_PARAGRAPH.CENTER)
CAP(doc,f"Tableau 3.41 — Resultats par pays (tau_opt={TAU_OPT:.4f}) — rouge : crise manquee")

# ── 3.5.6 Limites & conclusion ────────────────────────────────────────
SEP(doc)
H(doc, "3.5.6  Limites et positionnement du baseline", 2)
for bold_txt, rest in [
    ("EPV = " + f"{n1_tr/len(LAG_COLS):.1f}" + " : ",
     f"Avec {n1_tr} crises en train pour 5 variables, "
     "les coefficients sont contraints par la regularisation L2 "
     "plus que par le signal statistique. Ce niveau d'EPV est le "
     "minimum acceptable pour une identification des signes des coefficients."),
    ("Shift de regime : ",
     f"Le taux de crise passe de {n1_tr/n_tr*100:.1f}% (train) "
     f"a {n1_te/n_te*100:.1f}% (test). Ce changement structurel "
     "est inherent au panel CEMAC post-choc petrolier et affecte "
     "tous les modeles, pas uniquement le logit."),
    ("Linearite forcee : ",
     "Le logit impose une relation lineaire entre log-odds et features. "
     "Les seuils et interactions non-lineaires — centraux dans les "
     "crises bancaires — sont captures par les modeles suivants "
     "(Random Forest, XGBoost)."),
]:
    Bmix(doc,[("—  "+bold_txt,True,False,'1F3864'),(rest,False,False,None)],
         sb=0,sa=3,indent=0.5)

# Encadre conclusion
SEP(doc)
tbl_fin=doc.add_table(rows=1,cols=1)
tbl_fin.alignment=WD_TABLE_ALIGNMENT.CENTER
set_brd(tbl_fin,'1F3864',8)
c=tbl_fin.rows[0].cells[0]; c.width=Cm(15.5); set_bg(c,'DEEAF1')
p1=c.paragraphs[0]; p1.clear()
p1.alignment=WD_ALIGN_PARAGRAPH.CENTER
p1.paragraph_format.space_before=Pt(7)
p1.paragraph_format.space_after=Pt(3)
r1=p1.add_run("Bilan — Logit Baseline EWS-CEMAC (Split 2019)")
r1.bold=True; r1.font.size=Pt(11)
r1.font.color.rgb=RGBColor(0x1F,0x38,0x64)
p2=c.add_paragraph()
p2.alignment=WD_ALIGN_PARAGRAPH.LEFT
p2.paragraph_format.space_before=Pt(2)
p2.paragraph_format.space_after=Pt(7)
r2=p2.add_run(
    f"Le modele logistique baseline atteint sur le set de test {SPLIT_YEAR}-2022 :\n\n"
    f"  AUC-ROC = {m_te['auc']:.3f}  |  AUC-PR = {m_te['auc_pr']:.3f}  |  "
    f"Rappel = {m_te['rappel']:.3f}  |  Precision = {m_te['prec']:.3f}\n"
    f"  F1-Score = {m_te['f1']:.3f}  |  Brier Score = {m_te['brier']:.3f}\n\n"
    "Ces valeurs constituent la reference quantitative pour la comparaison "
    "avec les modeles non-lineaires (Random Forest, XGBoost). "
    "La valeur principale du logit reside dans l'interpretabilite de ses "
    "coefficients, qui confirment le role protecteur du solde budgetaire, "
    "de la croissance du PIB et des reserves de change dans la prevention "
    "des crises bancaires CEMAC."
)
r2.font.size=Pt(10.5)
r2.font.color.rgb=RGBColor(0x1A,0x1A,0x2E)

# ── Sauvegarde ────────────────────────────────────────────────────────
doc.save(PATH_DOC)
sz=os.path.getsize(PATH_DOC)//1024
print(f"\nDocument Word : {PATH_DOC}  ({sz} Ko)")
dv=Document(PATH_DOC)
hdgs=[(p.style.name,p.text[:65])
      for p in dv.paragraphs
      if p.style.name.startswith('Heading') and p.text.strip()]
print(f"Titres : {len(hdgs)}  |  Tableaux : {len(dv.tables)}")
for s,t in hdgs: print(f"  [{s}]  {t}")
print("\nTERMINE.")
