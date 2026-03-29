"""
04_xgboost.py
==============
EWS-CEMAC — XGBoost Classifier + SHAP Feature Importance, Reference Split 2019

PURPOSE
-------
Trains and evaluates an XGBoost gradient-boosted tree classifier on the CEMAC
panel (6 countries × 2000–2023) using the 2019 reference split, then computes
SHAP (SHapley Additive exPlanations) values to explain variable contributions.
Also produces a three-way comparison table (Logit-L2, RF, XGBoost).

XGBOOST HYPERPARAMETER JUSTIFICATION
--------------------------------------
All parameters are chosen conservatively to prevent overfitting on a small panel
(N_train ≈ 78 observations, EPV ≈ 4.8).

  n_estimators = 300, learning_rate = 0.05
      Slow learning rate combined with early stopping (implicit via a fixed
      number of rounds). Industry guidance: shrink learning_rate and increase
      n_estimators proportionally (Chen & Guestrin, 2016, KDD).

  max_depth = 2
      Shallower than RF (depth 3) to compensate for the additive nature of
      boosting — later trees correct residuals from early trees, so deeper
      trees risk memorising noise more quickly in sequential boosting.

  min_child_weight = 5
      Analogous to RF's min_samples_leaf; requires ≥ 5 observations in any
      child node before a split is accepted. Critical for small panels.

  subsample = 0.8
      Each tree is trained on a random 80% subsample of training observations
      (stochastic gradient boosting), which adds regularisation and reduces
      variance (Friedman, 2002).

  reg_lambda = 1.0, reg_alpha = 0.1
      L2 (Ridge) and L1 (Lasso) regularisation on leaf weights. Combined with
      the other constraints, these keep individual tree contributions small.

  scale_pos_weight = N_negative / N_positive
      XGBoost's built-in class imbalance correction, equivalent to the
      class_weight='balanced' formula in scikit-learn:
          scale_pos_weight = N_0 / N_1 = (N_train − n_crises) / n_crises
      This is recomputed from the training set at runtime.

  eval_metric = 'aucpr'
      Optimises the Area Under the Precision-Recall Curve during training,
      which is more informative than AUC-ROC for imbalanced datasets.

SHAP VALUES
-----------
SHAP (Lundberg & Lee, 2017, NeurIPS) provides a unified framework for
explaining the output of any machine learning model.  For a prediction f(x),
the SHAP value φ_j for feature j satisfies:

    f(x) = E[f(X)] + Σ_j φ_j(x)

Each φ_j measures the marginal contribution of feature j to the prediction
relative to the expected prediction.  Tree SHAP (Lundberg et al., 2020)
computes exact SHAP values for tree-based models in polynomial time.

IMPUTATION & SCALING: identical to scripts 02 and 03 (Bulmer rule, train-only).

AUTHORS
-------
  Françoise NGOUFACK, Pamphile MEZUI-MBENG, Samba NDIAYE — 2026
  Paper: "Do Early Warning Systems Survive Structural Breaks?
          Macroprudential Evidence from the CEMAC Monetary Union"
  Journal of Financial Stability [under review]
"""

import sys
sys.stdout.reconfigure(encoding='utf-8', errors='replace')

import os
import pandas as pd
import numpy as np
import warnings
warnings.filterwarnings('ignore')

import xgboost as xgb
import shap
from sklearn.linear_model import LogisticRegression
from sklearn.ensemble import RandomForestClassifier
from sklearn.inspection import permutation_importance
from sklearn.metrics import (
    roc_auc_score, roc_curve,
    precision_recall_curve, average_precision_score,
    confusion_matrix, brier_score_loss,
    precision_score, recall_score, f1_score
)
from sklearn.preprocessing import StandardScaler
from sklearn.impute import SimpleImputer
from sklearn.utils.class_weight import compute_class_weight

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.colors import LinearSegmentedColormap

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ══════════════════════════════════════════════════════════════════════
# PARAMETRES
# ══════════════════════════════════════════════════════════════════════
DATA_DIR = r"C:\Users\fngou\Desktop\Donnees_Memoire ML"
if not os.path.exists(DATA_DIR):
    DATA_DIR = r"C:\Users\fngou\Desktop\Données_Mémoire ML"

OUT_DIR    = r"C:\Users\fngou\Desktop\Chapitres du Memoire ML"
PATH_MACRO = DATA_DIR + r"\Dataset_Macro_CEMAC.csv"
PATH_DOC   = OUT_DIR  + r"\Section_37_XGB_StressVuln_2019.docx"
PATH_FIG1  = OUT_DIR  + r"\fig_xgb2019_roc_pr.png"
PATH_FIG2  = OUT_DIR  + r"\fig_xgb2019_confusion.png"
PATH_FIG3  = OUT_DIR  + r"\fig_xgb2019_shap_bar.png"
PATH_FIG4  = OUT_DIR  + r"\fig_xgb2019_shap_beeswarm.png"
PATH_FIG5  = OUT_DIR  + r"\fig_xgb2019_compare3.png"
PATH_FIG6  = OUT_DIR  + r"\fig_xgb2019_proba.png"

TARGET     = 'StressScore'
TARGET_BIN = 'Target_Stress2'
SPLIT_YEAR = 2019
SKEW_THR   = 1.0

SELECTED = [
    ('M2_croissance_pct',        'Croissance M2 (%)',        'C1',  0.69),
    ('Solde_budgetaire_pct_PIB', 'Solde budgetaire (% PIB)', 'C2',  0.99),
    ('PIB_croissance_reel_pct',  'Croissance PIB reel (%)',  'C3',  5.23),
    ('Reserves_USD',             'Reserves change (USD)',    'C4',  1.11),
    ('Rentes_petrole_pct_PIB',   'Rentes petrole (% PIB)',   'C5',  0.80),
]
BASE_COLS = [s[0] for s in SELECTED]
LAG_COLS  = [s[0]+'_lag1' for s in SELECTED]
LABELS    = {s[0]: s[1] for s in SELECTED}
LABELS_L  = {s[0]+'_lag1': s[1]+' (lag-1)' for s in SELECTED}
IMPUTE    = {s[0]: ('median' if abs(s[3]) > SKEW_THR else 'mean') for s in SELECTED}

BLUE='#1F3864'; LBLUE='#2E74B5'; ORANGE='#C0392B'
GRAY='#95A5A6'; GREEN='#1A7A2E'; TEAL='#148F77'; PURPLE='#7D3C98'

# ══════════════════════════════════════════════════════════════════════
# 1. PREPROCESSING (identique Logit & RF)
# ══════════════════════════════════════════════════════════════════════
print("="*66)
print("  EWS-CEMAC  |  XGBoost  |  StressScore >= 2  |  Split 2019")
print("="*66)

df = (pd.read_csv(PATH_MACRO)
        .sort_values(['Country','Year'])
        .reset_index(drop=True))
df[TARGET_BIN] = (df[TARGET] >= 2).astype(int)

# Lags sur panel complet avant split (variables predeterminees t-1 : pas de leakage)
for col in BASE_COLS:
    df[col+'_lag1'] = df.groupby('Country')[col].shift(1)

train_df = df[df['Year'] <  SPLIT_YEAR].copy().reset_index(drop=True)
test_df  = df[df['Year'] >= SPLIT_YEAR].copy().reset_index(drop=True)

X_tr_raw = train_df[LAG_COLS].copy()
X_te_raw = test_df[LAG_COLS].copy()
y_train  = train_df[TARGET_BIN].astype(int).values
y_test   = test_df[TARGET_BIN].astype(int).values
meta_te  = test_df[['Country','Year']].copy()

n_tr, n1_tr = len(y_train), int(y_train.sum())
n_te, n1_te = len(y_test),  int(y_test.sum())
n0_tr = n_tr - n1_tr

# Imputation
mean_cols   = [c+'_lag1' for c in BASE_COLS if IMPUTE[c]=='mean']
median_cols = [c+'_lag1' for c in BASE_COLS if IMPUTE[c]=='median']
imp_mean = SimpleImputer(strategy='mean');   imp_median = SimpleImputer(strategy='median')
if mean_cols:   imp_mean.fit(X_tr_raw[mean_cols])
if median_cols: imp_median.fit(X_tr_raw[median_cols])
for X in [X_tr_raw, X_te_raw]:
    if mean_cols:   X[mean_cols]   = imp_mean.transform(X[mean_cols])
    if median_cols: X[median_cols] = imp_median.transform(X[median_cols])
assert X_tr_raw.isna().sum().sum() == 0

# Normalisation
scaler  = StandardScaler()
X_train = pd.DataFrame(scaler.fit_transform(X_tr_raw), columns=LAG_COLS)
X_test  = pd.DataFrame(scaler.transform(X_te_raw),     columns=LAG_COLS)

# Class weights
cw_arr  = compute_class_weight('balanced', classes=np.array([0,1]), y=y_train)
cw_dict = {0: round(cw_arr[0],4), 1: round(cw_arr[1],4)}
sw      = np.where(y_train==1, cw_dict[1], cw_dict[0])
spw     = round(n0_tr / n1_tr, 2)   # scale_pos_weight natif XGBoost

print(f"\n  TRAIN : 2000-{SPLIT_YEAR-1} | {n_tr} obs | {n1_tr} vulnerables ({n1_tr/n_tr*100:.1f}%)")
print(f"  TEST  : {SPLIT_YEAR}-2023   | {n_te} obs | {n1_te} vulnerables ({n1_te/n_te*100:.1f}%)")
print(f"  EPV   : {n1_tr}/{len(LAG_COLS)} = {n1_tr/len(LAG_COLS):.1f}")
print(f"  scale_pos_weight : {n0_tr}/{n1_tr} = {spw}")

# ══════════════════════════════════════════════════════════════════════
# 2. HYPERPARAMETRES XGBoost
# ══════════════════════════════════════════════════════════════════════
XGB_PARAMS = dict(
    n_estimators      = 300,    # convergence stable avec lr=0.05
    max_depth         = 2,      # arbres-souches -> regularisation max
    learning_rate     = 0.05,   # apprentissage lent -> meilleure generalisation
    subsample         = 0.8,    # bagging stochastique
    colsample_bytree  = 1.0,    # 5 variables -> toutes utilisees
    min_child_weight  = 5,      # poids min par feuille ~ min_samples_leaf RF
    scale_pos_weight  = spw,    # correction desequilibre natif XGB
    reg_lambda        = 1.0,    # L2 regularisation (Ridge)
    reg_alpha         = 0.1,    # L1 regularisation (Lasso)
    objective         = 'binary:logistic',
    eval_metric       = 'aucpr',
    tree_method       = 'hist',
    random_state      = 42,
    n_jobs            = -1,
    verbosity         = 0,
)

# ══════════════════════════════════════════════════════════════════════
# 3. ENTRAINEMENT XGBOOST
# ══════════════════════════════════════════════════════════════════════
print("\n" + "="*66)
print(f"  XGBOOST  n={XGB_PARAMS['n_estimators']}  depth={XGB_PARAMS['max_depth']}"
      f"  lr={XGB_PARAMS['learning_rate']}  spw={spw}")
print("="*66)

model_xgb = xgb.XGBClassifier(**XGB_PARAMS)
model_xgb.fit(X_train, y_train,
              sample_weight=sw,
              eval_set=[(X_train, y_train)],
              verbose=False)

prob_tr = model_xgb.predict_proba(X_train)[:,1]
prob_te = model_xgb.predict_proba(X_test)[:,1]

print(f"\n  P_hat train : [{prob_tr.min():.4f}, {prob_tr.max():.4f}]  moy={prob_tr.mean():.4f}")
print(f"  P_hat test  : [{prob_te.min():.4f}, {prob_te.max():.4f}]  moy={prob_te.mean():.4f}")

# Seuil Youden sur train
fpr_tr, tpr_tr, thr_tr = roc_curve(y_train, prob_tr)
j_arr   = tpr_tr + (1 - fpr_tr) - 1
idx_opt = np.argmax(j_arr)
TAU_OPT = float(thr_tr[idx_opt])
TAU_DEF = 0.5
print(f"  Seuil Youden (train) : tau_opt = {TAU_OPT:.4f}")

# ══════════════════════════════════════════════════════════════════════
# 4. METRIQUES
# ══════════════════════════════════════════════════════════════════════
def metriques(y_true, y_prob, tau, label):
    y_pred = (y_prob >= tau).astype(int)
    cm = confusion_matrix(y_true, y_pred)
    tn,fp,fn,tp = (cm.ravel() if cm.size==4 else (cm[0,0],0,0,0)
                   if y_true.sum()==0 else (0,0,0,cm[0,0]))
    has = len(np.unique(y_true)) > 1
    auc    = roc_auc_score(y_true, y_prob)           if has else np.nan
    auc_pr = average_precision_score(y_true, y_prob) if has else np.nan
    rappel = recall_score(y_true, y_pred, zero_division=0)
    prec   = precision_score(y_true, y_pred, zero_division=0)
    f1     = f1_score(y_true, y_pred, zero_division=0)
    brier  = brier_score_loss(y_true, y_prob)
    spec   = tn/(tn+fp) if (tn+fp)>0 else 0.
    j_val  = rappel + spec - 1
    far    = fp/(fp+tn) if (fp+tn)>0 else 0.
    print(f"\n  --- {label} ---")
    print(f"  AUC-ROC : {auc:.4f}  |  AUC-PR : {auc_pr:.4f}")
    print(f"  Rappel  : {rappel:.4f} ({tp}/{tp+fn})  |  Precision : {prec:.4f}")
    print(f"  F1      : {f1:.4f}  |  Brier : {brier:.4f}")
    print(f"  J       : {j_val:.4f}  |  Spec={spec:.4f}  FAR={far:.4f}")
    print(f"  CM : TN={tn} FP={fp} FN={fn} TP={tp}")
    return dict(label=label, tau=tau, auc=auc, auc_pr=auc_pr,
                rappel=rappel, prec=prec, f1=f1, brier=brier,
                spec=spec, j=j_val, far=far,
                tn=int(tn), fp=int(fp), fn=int(fn), tp=int(tp))

print("\n" + "="*66)
print("  RESULTATS XGBOOST")
print("="*66)

m_tr  = metriques(y_train, prob_tr, TAU_OPT, f'TRAIN  tau={TAU_OPT:.4f}')
m_te  = metriques(y_test,  prob_te, TAU_OPT, f'TEST   tau={TAU_OPT:.4f} (calibre train)')
m_te5 = metriques(y_test,  prob_te, TAU_DEF, f'TEST   tau=0.50 (reference)')

# ── Importance XGBoost (gain, cover, weight) ──────────────────────────
score_types = ['gain','cover','weight']
imp_data = {}
for st in score_types:
    sc = model_xgb.get_booster().get_score(importance_type=st)
    imp_data[st] = {f: sc.get(f, 0) for f in LAG_COLS}

imp_df = pd.DataFrame(imp_data, index=LAG_COLS)
imp_df.index = [LABELS_L[c] for c in LAG_COLS]
imp_df = imp_df.div(imp_df.sum()).round(4)   # normaliser -> somme = 1
imp_df_sorted = imp_df.sort_values('gain', ascending=False)

print("\n  Importance XGBoost (normalisee) :")
print(f"  {'Variable':35s}  {'Gain':>8s}  {'Cover':>8s}  {'Weight':>8s}")
print("  " + "-"*63)
for idx, row in imp_df_sorted.iterrows():
    bar = chr(9608) * int(row['gain']*30)
    print(f"  {idx:35s}  {row['gain']:8.4f}  {row['cover']:8.4f}  {row['weight']:8.4f}  {bar}")

# ── SHAP values ──────────────────────────────────────────────────────
print("\n  Calcul des valeurs SHAP (TreeExplainer)...")
explainer   = shap.TreeExplainer(model_xgb)
shap_test   = explainer(X_test)
shap_train  = explainer(X_train)

# Remplacer noms de colonnes par labels lisibles
feat_names = [LABELS_L[c] for c in LAG_COLS]
shap_test.feature_names  = feat_names
shap_train.feature_names = feat_names

shap_vals_te = shap_test.values      # (n_test, 5)
shap_vals_tr = shap_train.values     # (n_train, 5)

print("  SHAP mean |value| par variable (test) :")
mean_abs = np.abs(shap_vals_te).mean(axis=0)
for i, (fn, mv) in enumerate(zip(feat_names, mean_abs)):
    bar = chr(9608) * int(mv * 30)
    print(f"    {fn:35s}  {mv:.4f}  {bar}")

# Analyse par pays (test)
print("\n  Analyse par pays — TEST :")
country_res = []
for country in sorted(meta_te['Country'].unique()):
    mask = meta_te['Country'].values == country
    y_c  = y_test[mask]; p_c = prob_te[mask]
    n_c  = len(y_c); nc1 = int(y_c.sum())
    pred = (p_c >= TAU_OPT).astype(int)
    tp_c = int(((y_c==1)&(pred==1)).sum())
    fn_c = int(((y_c==1)&(pred==0)).sum())
    fp_c = int(((y_c==0)&(pred==1)).sum())
    tn_c = int(((y_c==0)&(pred==0)).sum())
    auc_c = roc_auc_score(y_c, p_c) if len(np.unique(y_c))==2 else np.nan
    country_res.append((country,n_c,nc1,p_c.mean(),tp_c,fn_c,fp_c,tn_c,auc_c))
    auc_s = f"{auc_c:.3f}" if not np.isnan(auc_c) else 'n/a'
    print(f"    {country:30s}  N={n_c} Vuln={nc1}  "
          f"p_moy={p_c.mean():.3f}  TP={tp_c} FN={fn_c} FP={fp_c}  AUC={auc_s}")

# ── Re-calculer Logit & RF pour comparaison 3 modeles ─────────────────
print("\n  Recalcul Logit L2 et Random Forest pour comparaison...")

logit = LogisticRegression(penalty='l2', C=0.1, class_weight='balanced',
                            solver='lbfgs', max_iter=2000, random_state=42)
logit.fit(X_train, y_train)
pl_te = logit.predict_proba(X_test)[:,1]
pl_tr = logit.predict_proba(X_train)[:,1]
fpr_l, tpr_l, thr_l = roc_curve(y_train, pl_tr)
tau_l = float(thr_l[np.argmax(tpr_l + (1-fpr_l) - 1)])

rf = RandomForestClassifier(n_estimators=500, max_depth=3, min_samples_leaf=5,
                             max_features='sqrt', class_weight='balanced',
                             bootstrap=True, oob_score=True,
                             random_state=42, n_jobs=-1)
rf.fit(X_train, y_train)
pr_te = rf.predict_proba(X_test)[:,1]
pr_tr = rf.predict_proba(X_train)[:,1]
fpr_r, tpr_r, thr_r = roc_curve(y_train, pr_tr)
tau_r = float(thr_r[np.argmax(tpr_r + (1-fpr_r) - 1)])

def m_simple(y_true, y_prob, tau):
    y_pred = (y_prob >= tau).astype(int)
    cm = confusion_matrix(y_true, y_pred)
    tn,fp,fn,tp = cm.ravel() if cm.size==4 else (0,0,0,0)
    has = len(np.unique(y_true)) > 1
    return dict(
        auc    = roc_auc_score(y_true, y_prob)           if has else np.nan,
        auc_pr = average_precision_score(y_true, y_prob) if has else np.nan,
        rappel = recall_score(y_true, y_pred, zero_division=0),
        prec   = precision_score(y_true, y_pred, zero_division=0),
        f1     = f1_score(y_true, y_pred, zero_division=0),
        brier  = brier_score_loss(y_true, y_prob),
        tn=int(tn), fp=int(fp), fn=int(fn), tp=int(tp),
        far    = fp/(fp+tn) if (fp+tn)>0 else 0.,
        j      = recall_score(y_true,y_pred,zero_division=0)
                 + (tn/(tn+fp) if (tn+fp)>0 else 0.) - 1,
    )

ml = m_simple(y_test, pl_te, tau_l)
mr = m_simple(y_test, pr_te, tau_r)

print("\n" + "="*66)
print("  COMPARAISON 3 MODELES (test, tau_opt de chaque modele)")
print("="*66)
print(f"  {'Metrique':15s}  {'Logit L2':>10s}  {'Rand. Forest':>12s}  {'XGBoost':>10s}  {'Best':>6s}")
print("  " + "-"*62)
for k in ['auc','auc_pr','rappel','prec','f1','brier']:
    vl = ml[k]; vr = mr[k]; vx = m_te[k]
    vals = [vl, vr, vx]
    if k == 'brier':
        best_idx = int(np.argmin(vals))
    else:
        best_idx = int(np.argmax(vals))
    best_nm = ['Logit','RF','XGB'][best_idx]
    print(f"  {k:15s}  {vl:10.4f}  {vr:12.4f}  {vx:10.4f}  {best_nm:>6s}")

# ══════════════════════════════════════════════════════════════════════
# 5. FIGURES
# ══════════════════════════════════════════════════════════════════════
print("\nGeneration des figures...")

# ── Fig 1 : ROC + PR (3 modeles) ─────────────────────────────────────
fig, axes = plt.subplots(1, 2, figsize=(13, 5.5))
fig.suptitle(f"XGBoost vs RF vs Logit — Vulnerabilite CEMAC | Split {SPLIT_YEAR}",
             fontsize=12, fontweight='bold', color=BLUE)

fpr_x_te, tpr_x_te, _ = roc_curve(y_test,  prob_te)
fpr_x_tr, tpr_x_tr, _ = roc_curve(y_train, prob_tr)
fpr_r_te, tpr_r_te, _ = roc_curve(y_test,  pr_te)
fpr_l_te, tpr_l_te, _ = roc_curve(y_test,  pl_te)

ax = axes[0]
ax.fill_between(fpr_x_te, tpr_x_te, alpha=0.12, color=PURPLE)
ax.plot(fpr_x_te, tpr_x_te, color=PURPLE, lw=2.5,
        label=f"XGB   AUC={m_te['auc']:.3f}")
ax.plot(fpr_r_te, tpr_r_te, color=TEAL,   lw=1.8, ls='--',
        label=f"RF    AUC={mr['auc']:.3f}")
ax.plot(fpr_l_te, tpr_l_te, color=LBLUE,  lw=1.5, ls=':',
        label=f"Logit AUC={ml['auc']:.3f}")
ax.plot([0,1],[0,1], color=GRAY, lw=1, ls=':')
ax.scatter([m_te['far']], [m_te['rappel']], s=150, color=ORANGE, zorder=6,
           label=f"XGB tau_opt={TAU_OPT:.3f}  J={m_te['j']:.3f}")
ax.set_xlabel('FAR (1 - Specificite)', fontsize=11)
ax.set_ylabel('Rappel', fontsize=11)
ax.set_title('Courbe ROC', fontsize=11, fontweight='bold', color=BLUE)
ax.legend(fontsize=9); ax.grid(alpha=0.3)
ax.set_xlim([-0.02,1.02]); ax.set_ylim([-0.02,1.02])

prec_x_te, rec_x_te, _ = precision_recall_curve(y_test, prob_te)
prec_r_te, rec_r_te, _ = precision_recall_curve(y_test, pr_te)
prec_l_te, rec_l_te, _ = precision_recall_curve(y_test, pl_te)
ax = axes[1]
ax.fill_between(rec_x_te, prec_x_te, alpha=0.12, color=PURPLE)
ax.plot(rec_x_te, prec_x_te, color=PURPLE, lw=2.5,
        label=f"XGB   AUC-PR={m_te['auc_pr']:.3f}")
ax.plot(rec_r_te, prec_r_te, color=TEAL,   lw=1.8, ls='--',
        label=f"RF    AUC-PR={mr['auc_pr']:.3f}")
ax.plot(rec_l_te, prec_l_te, color=LBLUE,  lw=1.5, ls=':',
        label=f"Logit AUC-PR={ml['auc_pr']:.3f}")
ax.axhline(y_test.mean(), color=GRAY, lw=1, ls=':',
           label=f"Baseline ({y_test.mean():.2f})")
ax.scatter([m_te['rappel']], [m_te['prec']], s=150, color=ORANGE, zorder=6)
ax.set_xlabel('Rappel', fontsize=11); ax.set_ylabel('Precision', fontsize=11)
ax.set_title('Courbe Precision-Rappel', fontsize=11, fontweight='bold', color=BLUE)
ax.legend(fontsize=9); ax.grid(alpha=0.3)
ax.set_xlim([-0.02,1.02]); ax.set_ylim([-0.02,1.02])
plt.tight_layout()
plt.savefig(PATH_FIG1, dpi=150, bbox_inches='tight'); plt.close()
print("  Fig 1 (ROC+PR 3 modeles) : OK")

# ── Fig 2 : Confusion matrices ────────────────────────────────────────
fig, axes = plt.subplots(1, 2, figsize=(11, 4.5))
fig.suptitle(f"Matrices de confusion XGBoost — Test {SPLIT_YEAR}-2023",
             fontsize=12, fontweight='bold', color=BLUE)

def plot_cm(ax, m, title, col='#7D3C98'):
    cm_m = np.array([[m['tn'],m['fp']],[m['fn'],m['tp']]])
    cmap = LinearSegmentedColormap.from_list('c',['#FFFFFF', col])
    ax.imshow(cm_m, cmap=cmap, vmin=0, vmax=max(cm_m.max(),1))
    for i in range(2):
        for j in range(2):
            v = cm_m[i,j]
            c = 'white' if v > cm_m.max()*0.55 else 'black'
            ax.text(j,i,f"{[['TN','FP'],['FN','TP']][i][j]}\n{v}",
                    ha='center',va='center',fontsize=14,fontweight='bold',color=c)
    ax.set_xticks([0,1]); ax.set_yticks([0,1])
    ax.set_xticklabels(['Predit 0\n(Stable)','Predit 1\n(Vulnerable)'],fontsize=10)
    ax.set_yticklabels(['Reel 0\n(Stable)','Reel 1\n(Vulnerable)'],fontsize=10)
    ax.set_title(title, fontsize=11, fontweight='bold', color=BLUE)
    ax.set_xlabel(f"Rappel={m['rappel']:.2f}  Prec={m['prec']:.2f}  "
                  f"F1={m['f1']:.2f}  J={m['j']:.2f}", fontsize=9.5)

plot_cm(axes[0], m_te5, "Seuil par defaut (tau = 0.50)")
plot_cm(axes[1], m_te,  f"Seuil optimal Youden (tau = {TAU_OPT:.4f})")
plt.tight_layout()
plt.savefig(PATH_FIG2, dpi=150, bbox_inches='tight'); plt.close()
print("  Fig 2 (CM) : OK")

# ── Fig 3 : SHAP bar (importance globale) ─────────────────────────────
fig, ax = plt.subplots(figsize=(10, 4.5))
mean_abs_shap = np.abs(shap_vals_te).mean(axis=0)
order = np.argsort(mean_abs_shap)
colors_shap = [PURPLE if v > 0 else GRAY for v in mean_abs_shap[order]]
ax.barh([feat_names[i] for i in order],
        mean_abs_shap[order],
        color=PURPLE, edgecolor='white', height=0.6)
for i, (idx, v) in enumerate(zip(order, mean_abs_shap[order])):
    ax.text(v + 0.002, i, f'{v:.4f}', va='center', ha='left', fontsize=10)
ax.set_xlabel('Importance SHAP moyenne |E[phi]| (test)', fontsize=11)
ax.set_title('Importance SHAP globale — XGBoost (test 2019-2023)',
             fontsize=12, fontweight='bold', color=BLUE)
ax.grid(axis='x', alpha=0.3)
ax.tick_params(axis='y', labelsize=10)
plt.tight_layout()
plt.savefig(PATH_FIG3, dpi=150, bbox_inches='tight'); plt.close()
print("  Fig 3 (SHAP bar) : OK")

# ── Fig 4 : SHAP beeswarm ─────────────────────────────────────────────
fig, ax = plt.subplots(figsize=(11, 5))
order_bees = np.argsort(mean_abs_shap)[::-1]  # decroissant

# Beeswarm manuel
for i_feat, feat_idx in enumerate(order_bees):
    sv   = shap_vals_te[:, feat_idx]
    fval = X_test.iloc[:, feat_idx].values
    fval_norm = (fval - fval.min()) / (fval.max() - fval.min() + 1e-9)
    # jitter vertical
    jitter = np.random.RandomState(feat_idx).uniform(-0.2, 0.2, len(sv))
    colors_pts = plt.cm.RdBu_r(fval_norm)
    scatter = ax.scatter(sv, np.full(len(sv), i_feat) + jitter,
                         c=fval_norm, cmap='RdBu_r', vmin=0, vmax=1,
                         s=30, alpha=0.7, zorder=5)

ax.axvline(0, color='black', lw=0.8)
ax.set_yticks(range(len(feat_names)))
ax.set_yticklabels([feat_names[i] for i in order_bees], fontsize=10)
ax.set_xlabel('Valeur SHAP — impact sur P(Vulnerabilite)', fontsize=11)
ax.set_title('Diagramme Beeswarm SHAP — XGBoost (test 2019-2023)',
             fontsize=12, fontweight='bold', color=BLUE)
ax.grid(axis='x', alpha=0.3)
cbar = fig.colorbar(scatter, ax=ax, fraction=0.03, pad=0.02)
cbar.set_label('Valeur de la feature (bas -> haut)', fontsize=9)
cbar.set_ticks([0, 0.5, 1]); cbar.set_ticklabels(['Basse','Moyenne','Haute'])
plt.tight_layout()
plt.savefig(PATH_FIG4, dpi=150, bbox_inches='tight'); plt.close()
print("  Fig 4 (SHAP beeswarm) : OK")

# ── Fig 5 : Comparaison 3 modeles ─────────────────────────────────────
metrics_nm = ['AUC-ROC','AUC-PR','Rappel','Precision','F1-Score','Brier']
vals = {
    'Logit L2'     : [ml['auc'], ml['auc_pr'], ml['rappel'], ml['prec'], ml['f1'], ml['brier']],
    'Random Forest': [mr['auc'], mr['auc_pr'], mr['rappel'], mr['prec'], mr['f1'], mr['brier']],
    'XGBoost'      : [m_te['auc'], m_te['auc_pr'], m_te['rappel'], m_te['prec'], m_te['f1'], m_te['brier']],
}
colors_mod = [LBLUE, TEAL, PURPLE]

fig, ax = plt.subplots(figsize=(13, 5.5))
x = np.arange(len(metrics_nm)); w = 0.26
for i, (mod, col) in enumerate(zip(vals.keys(), colors_mod)):
    offset = (i - 1) * w
    bars = ax.bar(x + offset, vals[mod], w, label=mod,
                  color=col, edgecolor='white', alpha=0.88)
    for bar, v in zip(bars, vals[mod]):
        ax.text(bar.get_x()+bar.get_width()/2, bar.get_height()+0.01,
                f'{v:.3f}', ha='center', va='bottom', fontsize=8,
                color=col, fontweight='bold')

# Etoile sur le meilleur par metrique
for j, nm in enumerate(metrics_nm):
    row_vals = [vals[m][j] for m in vals]
    best_i   = int(np.argmin(row_vals) if nm == 'Brier' else np.argmax(row_vals))
    best_v   = row_vals[best_i]
    x_pos    = x[j] + (best_i - 1) * w
    ax.text(x_pos, best_v + 0.07, '★', ha='center', fontsize=11,
            color=colors_mod[best_i], fontweight='bold')

ax.set_xticks(x); ax.set_xticklabels(metrics_nm, fontsize=11)
ax.set_ylabel('Valeur (test set)', fontsize=11)
ax.set_ylim([0, 1.18])
ax.set_title(f"Comparaison 3 modeles — Test {SPLIT_YEAR}-2023 | StressScore >= 2 | ★ = meilleur",
             fontsize=12, fontweight='bold', color=BLUE)
ax.legend(fontsize=10, loc='upper right'); ax.grid(axis='y', alpha=0.3)
plt.tight_layout()
plt.savefig(PATH_FIG5, dpi=150, bbox_inches='tight'); plt.close()
print("  Fig 5 (comparaison 3 modeles) : OK")

# ── Fig 6 : Distribution P_hat ────────────────────────────────────────
fig, axes = plt.subplots(1, 2, figsize=(12, 4.5))
fig.suptitle("Distribution de P_hat(Vulnerabilite) — XGBoost",
             fontsize=12, fontweight='bold', color=BLUE)
for ax, (probs, y_true, title) in zip(axes, [
        (prob_tr, y_train, f"Train 2000-{SPLIT_YEAR-1}"),
        (prob_te, y_test,  f"Test {SPLIT_YEAR}-2023")]):
    ax.hist(probs[y_true==0], bins=20, alpha=0.6, color=LBLUE,
            label='Stable (y=0)', edgecolor='white')
    ax.hist(probs[y_true==1], bins=12, alpha=0.75, color=ORANGE,
            label='Vulnerable (y=1)', edgecolor='white')
    ax.axvline(TAU_OPT, color='red', ls='--', lw=2, label=f'tau_opt={TAU_OPT:.3f}')
    ax.axvline(TAU_DEF, color=GRAY, ls=':', lw=1.5, label='tau=0.50')
    ax.set_xlabel('P_hat(Vulnerabilite)', fontsize=11)
    ax.set_ylabel('Observations', fontsize=11)
    ax.set_title(title, fontsize=11, fontweight='bold', color=BLUE)
    ax.legend(fontsize=9.5); ax.grid(alpha=0.3)
plt.tight_layout()
plt.savefig(PATH_FIG6, dpi=150, bbox_inches='tight'); plt.close()
print("  Fig 6 (proba) : OK")

# ══════════════════════════════════════════════════════════════════════
# 6. DOCUMENT WORD
# ══════════════════════════════════════════════════════════════════════
print("\nConstruction du document Word...")

# Helpers (inlines)
def set_bg(cell,hx):
    tcPr=cell._tc.get_or_add_tcPr(); s=OxmlElement('w:shd')
    s.set(qn('w:val'),'clear'); s.set(qn('w:color'),'auto')
    s.set(qn('w:fill'),hx); tcPr.append(s)

def set_brd(t,color='1F3864',sz=8):
    tbl=t._tbl; tblPr=tbl.tblPr or OxmlElement('w:tblPr')
    b=OxmlElement('w:tblBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        el=OxmlElement(f'w:{side}')
        el.set(qn('w:val'),'single')
        el.set(qn('w:sz'),str(sz if side in('top','left','bottom','right') else 4))
        el.set(qn('w:space'),'0'); el.set(qn('w:color'),color); b.append(el)
    tblPr.append(b)

def cw(cell,text,bold=False,italic=False,sz=9.5,color=None,
       al=WD_ALIGN_PARAGRAPH.LEFT,sb=2,sa=2):
    p=cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    p.clear(); p.alignment=al
    p.paragraph_format.space_before=Pt(sb); p.paragraph_format.space_after=Pt(sa)
    cell.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
    r=p.add_run(text); r.bold=bold; r.italic=italic; r.font.size=Pt(sz)
    if color: r.font.color.rgb=RGBColor(*bytes.fromhex(color))

def H(doc,text,level,color='1F3864',sz=None):
    p=doc.add_paragraph(style=f'Heading {level}')
    r=p.add_run(text); r.bold=True
    r.font.size=Pt(sz or {1:13,2:12,3:11}[level])
    r.font.color.rgb=RGBColor(*bytes.fromhex(color))

def B(doc,text,italic=False,sz=10.5,sb=0,sa=4,indent=None):
    p=doc.add_paragraph(style='Normal')
    p.paragraph_format.space_before=Pt(sb); p.paragraph_format.space_after=Pt(sa)
    if indent: p.paragraph_format.left_indent=Cm(indent)
    r=p.add_run(text); r.font.size=Pt(sz); r.italic=italic

def Bmix(doc,segs,sb=0,sa=4,indent=None):
    p=doc.add_paragraph(style='Normal')
    p.paragraph_format.space_before=Pt(sb); p.paragraph_format.space_after=Pt(sa)
    if indent: p.paragraph_format.left_indent=Cm(indent)
    for txt,bold,italic,col in segs:
        r=p.add_run(txt); r.font.size=Pt(10.5); r.bold=bold; r.italic=italic
        if col: r.font.color.rgb=RGBColor(*bytes.fromhex(col))

def IMG(doc,path,w=15.0,cap_text=None,num=None):
    if not os.path.exists(path):
        B(doc,f"[Figure manquante : {path}]",italic=True); return
    p=doc.add_paragraph(style='Normal')
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(2)
    p.add_run().add_picture(path,width=Cm(w))
    if cap_text:
        pc=doc.add_paragraph(style='Normal')
        pc.alignment=WD_ALIGN_PARAGRAPH.CENTER
        pc.paragraph_format.space_before=Pt(2); pc.paragraph_format.space_after=Pt(10)
        rc=pc.add_run(f"Figure 3.{num} — {cap_text}" if num else cap_text)
        rc.italic=True; rc.font.size=Pt(9.5)
        rc.font.color.rgb=RGBColor(0x44,0x44,0x44)

def CAP(doc,text):
    p=doc.add_paragraph(style='Normal')
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(8)
    r=p.add_run(text); r.italic=True; r.font.size=Pt(9.5)
    r.font.color.rgb=RGBColor(0x44,0x44,0x44)

def SEP(doc):
    p=doc.add_paragraph(style='Normal')
    p.paragraph_format.space_before=Pt(10); p.paragraph_format.space_after=Pt(4)
    pPr=p._element.get_or_add_pPr(); pBdr=OxmlElement('w:pBdr')
    top=OxmlElement('w:top')
    top.set(qn('w:val'),'single'); top.set(qn('w:sz'),'4')
    top.set(qn('w:space'),'1'); top.set(qn('w:color'),'1F3864')
    pBdr.append(top); pPr.append(pBdr)

def INFO(doc,title,body,bg_c='DEEAF1',brd_c='1F3864'):
    t=doc.add_table(rows=1,cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    set_brd(t,color=brd_c,sz=6)
    c=t.rows[0].cells[0]; c.width=Cm(15.5); set_bg(c,bg_c)
    p1=c.paragraphs[0]; p1.clear()
    p1.paragraph_format.space_before=Pt(5); p1.paragraph_format.space_after=Pt(2)
    r1=p1.add_run(title); r1.bold=True; r1.font.size=Pt(10)
    r1.font.color.rgb=RGBColor(*bytes.fromhex(brd_c))
    p2=c.add_paragraph()
    p2.paragraph_format.space_before=Pt(0); p2.paragraph_format.space_after=Pt(5)
    r2=p2.add_run(body); r2.font.size=Pt(10); r2.italic=True
    r2.font.color.rgb=RGBColor(0x1A,0x1A,0x2E)
    doc.add_paragraph(style='Normal').paragraph_format.space_after=Pt(6)

def FBOX(doc,text):
    t=doc.add_table(rows=1,cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    set_brd(t,color='AAAAAA',sz=4)
    c=t.rows[0].cells[0]; c.width=Cm(14); set_bg(c,'F2F2F2')
    p=c.paragraphs[0]; p.clear()
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before=Pt(5); p.paragraph_format.space_after=Pt(5)
    r=p.add_run(text); r.font.size=Pt(10.5)
    r.font.name='Courier New'; r.font.color.rgb=RGBColor(0x1F,0x38,0x64)
    doc.add_paragraph(style='Normal').paragraph_format.space_after=Pt(4)

# ─── Construction ────────────────────────────────────────────────────
doc = Document()
for s in doc.sections:
    s.top_margin=Cm(2.5); s.bottom_margin=Cm(2.5)
    s.left_margin=Cm(3.0); s.right_margin=Cm(2.5)

H(doc,"3.7  XGBoost — Regime de Vulnerabilite Macro-Financiere",1)
B(doc,
  "XGBoost (eXtreme Gradient Boosting) est le deuxieme modele non-lineaire "
  "de la strategie de comparaison. Contrairement au Random Forest "
  "(bagging parallele d'arbres independants), XGBoost construit "
  "sequentiellement des arbres qui corrigent les erreurs du modele precedent "
  "(boosting). Il integre nativement une regularisation L1+L2, "
  "un mecanisme de ponderation des classes (scale_pos_weight) et "
  "des valeurs SHAP pour l'interpretabilite — atouts majeurs "
  "pour un EWS publiable.",
  sb=6, sa=6)

# ── 3.7.1 Hyperparametres ─────────────────────────────────────────────
H(doc,"3.7.1  Specification et justification des hyperparametres",2)

HP_ROWS=[
    ('n_estimators','300',
     'Convergence stable avec learning_rate=0.05 (Chen & Guestrin, 2016)'),
    ('max_depth','2',
     'Arbres-souches (depth=2) : regularisation maximale pour EPV=4.8'),
    ('learning_rate','0.05',
     'Apprentissage lent -> meilleure generalisation hors echantillon'),
    ('subsample','0.8',
     'Bagging stochastique : 80% observations par arbre -> regularisation'),
    ('colsample_bytree','1.0',
     '5 variables -> toutes utilisees a chaque arbre'),
    ('min_child_weight','5',
     'Poids min par noeud terminal ~ min_samples_leaf=5 du RF'),
    ('scale_pos_weight',str(spw),
     f'n0/n1 = {n0_tr}/{n1_tr} : correction desequilibre classe natif XGB'),
    ('reg_lambda','1.0',
     'Regularisation L2 (Ridge) : controle la magnitude des scores feuilles'),
    ('reg_alpha','0.1',
     'Regularisation L1 (Lasso) : favorise la parcimonie des arbres'),
    ('eval_metric','aucpr',
     'Optimisation sur AUC-PR durant le boosting : mieux adapte aux classes desequilibrees'),
]
tbl_hp=doc.add_table(rows=len(HP_ROWS)+1,cols=3)
tbl_hp.style='Table Grid'; tbl_hp.alignment=WD_TABLE_ALIGNMENT.CENTER
set_brd(tbl_hp,'1F3864',8)
for row in tbl_hp.rows:
    for i,c in enumerate(row.cells): c.width=Cm([4.0,2.5,9.5][i])
for i,h_ in enumerate(['Hyperparametre','Valeur','Justification']):
    set_bg(tbl_hp.rows[0].cells[i],'1F3864')
    cw(tbl_hp.rows[0].cells[i],h_,bold=True,sz=9,color='FFFFFF',al=WD_ALIGN_PARAGRAPH.CENTER)
for r_i,(hp,val,just) in enumerate(HP_ROWS):
    row=tbl_hp.rows[r_i+1]
    bg_c='F0ECF8' if r_i%2==0 else 'FFFFFF'
    for i,v in enumerate([hp,val,just]):
        set_bg(row.cells[i],bg_c)
        cw(row.cells[i],v,sz=9,bold=(i==0),
           al=WD_ALIGN_PARAGRAPH.CENTER if i<2 else WD_ALIGN_PARAGRAPH.LEFT)
CAP(doc,"Tableau 3.47 — Hyperparametres XGBoost et justification")

FBOX(doc,
     "XGBoost :  P(Vuln_{t+1} | X_t) = sigma( sum_{m=1}^{300} f_m(X_t) )\n\n"
     "  f_m : arbre CART  depth=2  |  Optimisation : AUC-PR (gradient boosting)\n"
     f"  Reg : L1={XGB_PARAMS['reg_alpha']}  L2={XGB_PARAMS['reg_lambda']}  "
     f"subsample={XGB_PARAMS['subsample']}  lr={XGB_PARAMS['learning_rate']}\n"
     f"  scale_pos_weight = {spw}  |  tau_opt = {TAU_OPT:.4f} (Youden / train)")

# ── 3.7.2 Resultats ───────────────────────────────────────────────────
SEP(doc)
H(doc,"3.7.2  Resultats et metriques de performance",2)

MROWS=[
    ('AUC-ROC',    m_tr['auc'],    m_te5['auc'],    m_te['auc'],    False),
    ('AUC-PR',     m_tr['auc_pr'], m_te5['auc_pr'], m_te['auc_pr'], False),
    ('Rappel',     m_tr['rappel'], m_te5['rappel'],  m_te['rappel'],  False),
    ('Precision',  m_tr['prec'],   m_te5['prec'],   m_te['prec'],   False),
    ('F1-Score',   m_tr['f1'],     m_te5['f1'],     m_te['f1'],     False),
    ('Brier Score',m_tr['brier'],  m_te5['brier'],  m_te['brier'],  True),
]
KEY={'AUC-ROC','AUC-PR','Rappel','F1-Score'}

tbl_r=doc.add_table(rows=len(MROWS)+2,cols=4)
tbl_r.style='Table Grid'; tbl_r.alignment=WD_TABLE_ALIGNMENT.CENTER
set_brd(tbl_r,'1F3864',8)
for row in tbl_r.rows:
    for i,c in enumerate(row.cells): c.width=Cm([3.8,3.0,3.0,3.0][i])
for i,h_ in enumerate(['Metrique','Train','Test tau=0.50','Test tau_opt']):
    set_bg(tbl_r.rows[0].cells[i],['1F3864','7D3C98','1A5276','1A5276'][i])
    cw(tbl_r.rows[0].cells[i],h_,bold=True,sz=9,color='FFFFFF',al=WD_ALIGN_PARAGRAPH.CENTER)
set_bg(tbl_r.rows[1].cells[0],'2C3E50'); cw(tbl_r.rows[1].cells[0],'',sz=8)
for i,val in enumerate([f'tau={TAU_OPT:.4f}','tau=0.50',f'tau={TAU_OPT:.4f}']):
    set_bg(tbl_r.rows[1].cells[i+1],['7D3C98','1A5276','1A5276'][i])
    cw(tbl_r.rows[1].cells[i+1],val,sz=8,color='FFFFFF',al=WD_ALIGN_PARAGRAPH.CENTER)
for r_i,(name,v_tr,v_te5,v_te,low) in enumerate(MROWS):
    row=tbl_r.rows[r_i+2]
    bg0='F5EEF8' if name in KEY else ('EDF3FB' if r_i%2==0 else 'FFFFFF')
    set_bg(row.cells[0],bg0)
    cw(row.cells[0],name,bold=(name in KEY),sz=9.5)
    for ci,(val,is_opt) in enumerate([(v_tr,False),(v_te5,False),(v_te,True)]):
        better=(val>v_te5) if not low else (val<v_te5)
        bg_c=('E8F5E9' if is_opt and better else
              'FDECEA' if is_opt and not better else bg0)
        set_bg(row.cells[ci+1],bg_c)
        cw(row.cells[ci+1],f"{val:.4f}",bold=is_opt,sz=9.5,
           al=WD_ALIGN_PARAGRAPH.CENTER,
           color=('1A7A2E' if is_opt and better else
                  'C0392B' if is_opt and not better else None))
CAP(doc,f"Tableau 3.48 — Resultats XGBoost (StressScore>=2, split {SPLIT_YEAR}, "
        f"tau_opt={TAU_OPT:.4f})")

IMG(doc,PATH_FIG1,w=15.5,
    cap_text=f"Courbes ROC et PR — XGBoost vs RF vs Logit (test {SPLIT_YEAR}-2023)",
    num='20')
IMG(doc,PATH_FIG2,w=13.5,
    cap_text=f"Matrices de confusion XGBoost — Test {SPLIT_YEAR}-2023 (tau=0.50 vs tau_opt={TAU_OPT:.4f})",
    num='21')

# ── 3.7.3 Interpretabilite SHAP ───────────────────────────────────────
SEP(doc)
H(doc,"3.7.3  Interpretabilite par valeurs SHAP",2)
B(doc,
  "Les valeurs SHAP (SHapley Additive exPlanations, Lundberg & Lee, 2017) "
  "decompposent la prediction de chaque observation en contributions "
  "additives par variable. Elles satisfont les proprietes d'efficacite, "
  "symetrie et absence de joueur nul — ce qui en fait la methode d'interpretabilite "
  "la plus rigoureuse pour les modeles ensemblistes. "
  "Les valeurs sont calculees sur le set de test (TreeExplainer).",
  sb=0, sa=4)

# Tableau SHAP
shap_table_data = []
for i_f, fn in enumerate(feat_names):
    sv = shap_vals_te[:, i_f]
    shap_table_data.append((
        fn,
        f"{mean_abs_shap[i_f]:.4f}",
        f"{sv[y_test==1].mean():+.4f}",
        f"{sv[y_test==0].mean():+.4f}",
        "Amplifie la vulnerabilite" if sv[y_test==1].mean()>0 else "Attenue la vulnerabilite"
    ))
shap_table_data.sort(key=lambda x: -float(x[1]))

tbl_sh=doc.add_table(rows=len(shap_table_data)+1,cols=5)
tbl_sh.style='Table Grid'; tbl_sh.alignment=WD_TABLE_ALIGNMENT.CENTER
set_brd(tbl_sh,'1F3864',8)
for row in tbl_sh.rows:
    for i,c in enumerate(row.cells): c.width=Cm([4.0,2.0,2.2,2.2,5.6][i])
for i,h_ in enumerate(['Variable (lag-1)','|SHAP| moy.',
                        'SHAP moy. (vuln.=1)','SHAP moy. (stable=0)','Interpretation']):
    set_bg(tbl_sh.rows[0].cells[i],'1F3864')
    cw(tbl_sh.rows[0].cells[i],h_,bold=True,sz=9,color='FFFFFF',al=WD_ALIGN_PARAGRAPH.CENTER)
for r_i,(fn,abs_s,s1,s0,interp) in enumerate(shap_table_data):
    row=tbl_sh.rows[r_i+1]
    bg_c='F5EEF8' if r_i%2==0 else 'FFFFFF'
    is_risk = float(s1) > 0
    for i,val in enumerate([fn,abs_s,s1,s0,interp]):
        set_bg(row.cells[i],bg_c)
        cw(row.cells[i],val,sz=9,bold=(i in [1,2]),
           color=('C0392B' if i==2 and is_risk else
                  '1A5276' if i==2 and not is_risk else None),
           al=WD_ALIGN_PARAGRAPH.CENTER if i in [1,2,3] else WD_ALIGN_PARAGRAPH.LEFT)
CAP(doc,"Tableau 3.49 — Valeurs SHAP par variable (test 2019-2023) — "
        "rouge : contribution positive a P(Vulnerabilite)")

IMG(doc,PATH_FIG3,w=12.0,
    cap_text="Importance SHAP globale |E[phi]| — XGBoost (test 2019-2023)",num='22')
IMG(doc,PATH_FIG4,w=15.5,
    cap_text="Diagramme Beeswarm SHAP — distribution des contributions individuelles",num='23')

INFO(doc,
     "Interpretation SHAP pour BEAC/COBAC :",
     "Le diagramme beeswarm permet au superviseur d'identifier, "
     "pour chaque pays et chaque annee, quelle variable a le plus "
     "contribue a l'alerte ou a l'absence d'alerte. "
     "Cette interpretabilite locale est un atout operationnel majeur "
     "par rapport aux modeles boite noire purs.",
     bg_c='E8F8F0', brd_c='148F77')

# ── 3.7.4 Analyse par pays ────────────────────────────────────────────
SEP(doc)
H(doc,f"3.7.4  Analyse par pays — Test {SPLIT_YEAR}-2023",2)
tbl_p=doc.add_table(rows=len(country_res)+1,cols=7)
tbl_p.style='Table Grid'; tbl_p.alignment=WD_TABLE_ALIGNMENT.CENTER
set_brd(tbl_p,'1F3864',8)
for row in tbl_p.rows:
    for i,c in enumerate(row.cells): c.width=Cm([4.0,1.0,1.5,2.0,1.0,1.0,1.0][i])
for i,h_ in enumerate(['Pays','N','Vuln.','P_hat moy.','TP','FN','FP']):
    set_bg(tbl_p.rows[0].cells[i],'1F3864')
    cw(tbl_p.rows[0].cells[i],h_,bold=True,sz=9,color='FFFFFF',al=WD_ALIGN_PARAGRAPH.CENTER)
for r_i,(country,n_c,nc1,avg_p,tp_c,fn_c,fp_c,tn_c,auc_c) in enumerate(country_res):
    row=tbl_p.rows[r_i+1]
    rbg='FDECEA' if fn_c>0 else ('E8F5E9' if nc1>0 else 'EDF3FB')
    for i,val in enumerate([country,str(n_c),str(nc1),
                             f"{avg_p:.3f}",str(tp_c),str(fn_c),str(fp_c)]):
        set_bg(row.cells[i],rbg)
        cw(row.cells[i],val,sz=9.5,bold=(i==5 and fn_c>0),
           color=('C0392B' if fn_c>0 and i==5 else None),
           al=WD_ALIGN_PARAGRAPH.LEFT if i==0 else WD_ALIGN_PARAGRAPH.CENTER)
CAP(doc,f"Tableau 3.50 — Resultats XGB par pays (tau_opt={TAU_OPT:.4f}) — rouge : vulnerabilite manquee")

# ── 3.7.5 Comparaison 3 modeles ───────────────────────────────────────
SEP(doc)
H(doc,"3.7.5  Tableau de synthese — 3 modeles",2)
B(doc,
  "Le tableau suivant consolide les performances des trois modeles "
  "sur le set de test commun (2019-2023). Chaque modele est evalue "
  "a son seuil tau_opt propre (Youden sur train). "
  "L'etoile ★ indique le meilleur modele par metrique.",
  sb=0, sa=4)

SYNTH=[
    ('AUC-ROC',    ml['auc'],    mr['auc'],    m_te['auc'],    False),
    ('AUC-PR',     ml['auc_pr'], mr['auc_pr'], m_te['auc_pr'], False),
    ('Rappel',     ml['rappel'], mr['rappel'],  m_te['rappel'],  False),
    ('Precision',  ml['prec'],   mr['prec'],   m_te['prec'],   False),
    ('F1-Score',   ml['f1'],     mr['f1'],     m_te['f1'],     False),
    ('Brier Score',ml['brier'],  mr['brier'],  m_te['brier'],  True),
    ('tau_opt',    tau_l,        tau_r,        TAU_OPT,        None),
]

tbl_syn=doc.add_table(rows=len(SYNTH)+1,cols=4)
tbl_syn.style='Table Grid'; tbl_syn.alignment=WD_TABLE_ALIGNMENT.CENTER
set_brd(tbl_syn,'1F3864',8)
for row in tbl_syn.rows:
    for i,c in enumerate(row.cells): c.width=Cm([3.5,3.5,3.5,3.5][i])
for i,h_ in enumerate(['Metrique','Logit L2','Random Forest','XGBoost']):
    set_bg(tbl_syn.rows[0].cells[i],['1F3864','2E74B5','148F77','7D3C98'][i])
    cw(tbl_syn.rows[0].cells[i],h_,bold=True,sz=9.5,color='FFFFFF',al=WD_ALIGN_PARAGRAPH.CENTER)

for r_i,(name,vl,vr,vx,low) in enumerate(SYNTH):
    row=tbl_syn.rows[r_i+1]
    bg0='EDF3FB' if r_i%2==0 else 'FFFFFF'
    set_bg(row.cells[0],bg0); cw(row.cells[0],name,bold=True,sz=9.5)
    if low is None:
        for i,val in enumerate([vl,vr,vx]):
            set_bg(row.cells[i+1],bg0)
            cw(row.cells[i+1],f"{val:.4f}",sz=9.5,al=WD_ALIGN_PARAGRAPH.CENTER)
    else:
        vals3=[vl,vr,vx]
        best_i = int(np.argmin(vals3) if low else np.argmax(vals3))
        for i,val in enumerate([vl,vr,vx]):
            is_best = (i == best_i)
            bg_v = 'E8F5E9' if is_best else bg0
            set_bg(row.cells[i+1],bg_v)
            txt = f"{val:.4f}" + (" ★" if is_best else "")
            cw(row.cells[i+1],txt,bold=is_best,sz=9.5,
               al=WD_ALIGN_PARAGRAPH.CENTER,
               color=('1A7A2E' if is_best else None))
CAP(doc,"Tableau 3.51 — Synthese comparative : Logit L2 vs RF vs XGBoost (test tau_opt) — ★ meilleur")

IMG(doc,PATH_FIG5,w=15.5,
    cap_text=f"Comparaison des 6 metriques — Logit L2, Random Forest, XGBoost — Test {SPLIT_YEAR}-2023",
    num='24')

# ── 3.7.6 Limites & conclusion ────────────────────────────────────────
SEP(doc)
H(doc,"3.7.6  Discussion et conclusion de la comparaison",2)
for bold_txt, rest in [
    ("Shift de regime (probleme commun aux 3 modeles) : ",
     f"Les trois modeles sont entraines sur un regime de vulnerabilite "
     f"caracteristique de 2014-2018 (choc petrolier) et evalues sur "
     f"un regime post-2019 (Covid, sanctions, consolidation). "
     "Ce changement structurel est la principale limitation — "
     "il est inherent au panel CEMAC, pas au choix du modele."),
    ("Apport de XGBoost vs RF : ",
     "XGBoost offre une regularisation native L1+L2, une optimisation "
     "directe sur AUC-PR et des valeurs SHAP interpretables. "
     "Ces atouts sont decisifs pour la publication scientifique "
     "et l'acceptabilite par BEAC/COBAC."),
    ("Prochaine etape — enrichissement des features : ",
     "Les resultats justifient d'explorer des features supplementaires "
     "(variables en niveau, interactions, features d'alerte precoce "
     "comme le ratio credit/PIB ou l'ecart de croissance regional) "
     "ou un horizon de prediction different (h=2 ans)."),
]:
    Bmix(doc,[("—  "+bold_txt,True,False,'1F3864'),(rest,False,False,None)],
         sb=0,sa=3,indent=0.5)

# Bilan final
SEP(doc)
tbl_fin=doc.add_table(rows=1,cols=1)
tbl_fin.alignment=WD_TABLE_ALIGNMENT.CENTER
set_brd(tbl_fin,'1F3864',8)
c=tbl_fin.rows[0].cells[0]; c.width=Cm(15.5); set_bg(c,'F5EEF8')
p1=c.paragraphs[0]; p1.clear()
p1.alignment=WD_ALIGN_PARAGRAPH.CENTER
p1.paragraph_format.space_before=Pt(7); p1.paragraph_format.space_after=Pt(3)
r1=p1.add_run("Bilan Comparatif — 3 Modeles EWS-CEMAC | Vulnerabilite macro-financiere | Split 2019")
r1.bold=True; r1.font.size=Pt(11); r1.font.color.rgb=RGBColor(0x1F,0x38,0x64)
p2=c.add_paragraph()
p2.alignment=WD_ALIGN_PARAGRAPH.LEFT
p2.paragraph_format.space_before=Pt(3); p2.paragraph_format.space_after=Pt(7)

# Determine best model per metric
best_models = []
for nm, vl, vr, vx, low in SYNTH[:-1]:
    vals3=[vl,vr,vx]; nms=['Logit','RF','XGB']
    bi = int(np.argmin(vals3) if low else np.argmax(vals3))
    best_models.append(f"{nm} -> {nms[bi]} ({vals3[bi]:.3f})")

r2=p2.add_run(
    f"Cible : StressScore >= 2  |  Split : 2000-{SPLIT_YEAR-1} / {SPLIT_YEAR}-2023\n"
    f"Variables : 5 lag-1  |  Seuil : Youden sur train\n\n"
    "Meilleur modele par metrique :\n" +
    "\n".join(f"  {s}" for s in best_models) +
    f"\n\nConclusion : les trois modeles partagent la contrainte du shift de regime "
    f"({n1_tr/n_tr*100:.0f}% train -> {n1_te/n_te*100:.0f}% test). "
    f"XGBoost se distingue par l'interpretabilite SHAP et la regularisation L1+L2 "
    f"— criteres prioritaires pour la publication et le deploiement BEAC/COBAC."
)
r2.font.size=Pt(10.5); r2.font.color.rgb=RGBColor(0x1A,0x1A,0x2E)

IMG(doc,PATH_FIG6,w=15.5,
    cap_text=f"Distribution de P_hat(Vulnerabilite) — XGBoost (train et test)",num='25')

# Sauvegarde
doc.save(PATH_DOC)
sz=os.path.getsize(PATH_DOC)//1024
print(f"\nDocument Word : {PATH_DOC}  ({sz} Ko)")
dv=Document(PATH_DOC)
hdgs=[(p.style.name,p.text[:65])
      for p in dv.paragraphs
      if p.style.name.startswith('Heading') and p.text.strip()]
print(f"Titres : {len(hdgs)}  |  Tableaux : {len(dv.tables)}")
for s,t in hdgs: print(f"  [{s}]  {t}")
print("\nTERMINE.")
